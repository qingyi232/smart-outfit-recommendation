# -*- coding: utf-8 -*-
"""替换论文中代码截图为白色背景图片"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'thesis_images')
os.makedirs(IMG_DIR, exist_ok=True)

SRC = os.path.join(BASE_DIR, '通信工程（专转本）Z2024130143陈鎏-基于环境信息感知的智能出行穿衣服系统.docx')
OUT = SRC

CODE_BLOCKS = {
    '图5-1': '''def get_weather(city, target_date=None):
    """获取指定城市的天气数据（核心接口）"""
    if target_date is None:
        target_date = date.today()
    elif isinstance(target_date, str):
        target_date = datetime.strptime(target_date, '%Y-%m-%d').date()

    cached = WeatherRecord.query.filter_by(
        city=city, date=target_date).first()
    if cached and (datetime.utcnow()
            - cached.fetched_at).total_seconds() < 3600:
        return cached.to_dict()

    api_key = current_app.config.get('WEATHER_API_KEY', '')
    api_data = None
    if api_key and target_date == date.today():
        api_data = _fetch_from_qweather(city, api_key)

    weather_data = _generate_weather_for_city(city, target_date)
    if api_data:
        weather_data.update(api_data)
    return cached.to_dict()''',

    '图5-2': '''@auth_bp.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username', '').strip()
    password = data.get('password', '')

    user = User.query.filter(
        (User.username == username) |
        (User.email == username)
    ).first()

    if not user or not user.check_password(password):
        return jsonify(code=401, message='用户名或密码错误')
    if not user.is_active:
        return jsonify(code=403, message='账号已被禁用')

    token = create_access_token(
        identity=user.id,
        expires_delta=timedelta(days=7)
    )
    return jsonify(code=200,
        data={'token': token, 'user': user.to_dict()})''',

    '图5-5': '''const service = axios.create({
    baseURL: '/api',
    timeout: 10000
})

service.interceptors.request.use(config => {
    const token = localStorage.getItem('token')
    if (token) {
        config.headers.Authorization = `Bearer ${token}`
    }
    return config
})

service.interceptors.response.use(
    response => response.data,
    error => {
        if (error.response?.status === 401) {
            localStorage.clear()
            router.push('/login')
        }
        return Promise.reject(error)
    }
)''',

    '图5-10': '''class DecisionTreeRecommender:
    def __init__(self, training_data=None,
                 model_name='standard',
                 max_depth=8, min_samples_split=3):
        self.model_name = model_name
        self.model = DecisionTreeClassifier(
            max_depth=max_depth,
            min_samples_split=min_samples_split,
            min_samples_leaf=2, random_state=42)
        self._training_data = training_data or TRAINING_DATA
        self._train()

    def _train(self):
        data = np.array(self._training_data)
        X = data[:, :-1]
        y = data[:, -1]
        self.model.fit(X, y)

    def predict_warmth_level(self, temp, humidity,
            wind_speed, precipitation, aqi, scene):
        scene_id = SCENE_LIST.index(scene) \\
            if scene in SCENE_LIST else 0
        features = np.array([[temp, humidity, wind_speed,
                              precipitation, aqi, scene_id]])
        return int(self.model.predict(features)[0])''',

    '图5-11': '''def get_similar_users(self, user_id, top_n=5):
    """计算用户相似度，找到最相似的N个用户"""
    all_prefs = UserPreference.query.all()
    if not all_prefs:
        return []

    user_items = {}
    for p in all_prefs:
        user_items.setdefault(
            p.user_id, {})[p.clothing_id] = p.preference_score

    if user_id not in user_items:
        return []

    target = user_items[user_id]
    similarities = []
    for other_id, other_prefs in user_items.items():
        if other_id == user_id:
            continue
        common = set(target.keys()) & set(other_prefs.keys())
        if len(common) < 2:
            continue
        vec_a = np.array([target[i] for i in common])
        vec_b = np.array([other_prefs[i] for i in common])
        norm_a = np.linalg.norm(vec_a)
        norm_b = np.linalg.norm(vec_b)
        if norm_a == 0 or norm_b == 0:
            continue
        cos_sim = np.dot(vec_a, vec_b) / (norm_a * norm_b)
        similarities.append((other_id, cos_sim))
    return sorted(similarities,
        key=lambda x: x[1], reverse=True)[:top_n]''',
}


def generate_code_image(code_text, output_path):
    lines = code_text.strip().split('\n')
    line_height = 20
    padding = {'top': 15, 'bottom': 15, 'left': 60, 'right': 25}

    try:
        code_font = ImageFont.truetype('C:/Windows/Fonts/consola.ttf', 14)
        num_font = ImageFont.truetype('C:/Windows/Fonts/consola.ttf', 13)
    except OSError:
        code_font = ImageFont.load_default()
        num_font = code_font

    max_len = max(len(l) for l in lines)
    w = max(750, int(padding['left'] + max_len * 8.8 + padding['right']))
    h = int(padding['top'] + len(lines) * line_height + padding['bottom'])

    img = Image.new('RGB', (w, h), '#FFFFFF')
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w-1, h-1], outline='#CCCCCC', width=1)
    draw.line([(padding['left']-12, 1), (padding['left']-12, h-2)], fill='#E0E0E0')

    for i, line in enumerate(lines):
        y = padding['top'] + i * line_height
        draw.text((12, y), str(i+1).rjust(3), fill='#AAAAAA', font=num_font)
        draw.text((padding['left'], y), line, fill='#2E2E2E', font=code_font)

    img.save(output_path, dpi=(150, 150))


def has_drawing(para_element):
    xml = etree.tostring(para_element, encoding='unicode')
    return 'w:drawing' in xml or 'w:pict' in xml


def main():
    print('步骤1: 生成白色背景代码图片...')
    img_paths = {}
    for key, code in CODE_BLOCKS.items():
        fname = f'code_{key.replace("图", "").replace("-", "_")}.png'
        path = os.path.join(IMG_DIR, fname)
        generate_code_image(code, path)
        img_paths[key] = path
        print(f'  生成: {fname}')

    print('\n步骤2: 替换文档中的代码截图...')
    doc = Document(SRC)
    body = doc.element.body
    paragraphs = list(doc.paragraphs)

    caption_map = {}
    for i, p in enumerate(paragraphs):
        if p.style and p.style.name == '图目录':
            text = p.text.strip().replace(' ', '')
            for key in img_paths:
                key_clean = key.replace(' ', '')
                if text.startswith(key_clean):
                    caption_map[i] = key

    print(f'  找到 {len(caption_map)} 个代码图题')

    replaced = 0
    for cap_idx, key in caption_map.items():
        img_path = img_paths[key]
        caption_elem = paragraphs[cap_idx]._element

        for offset in range(1, 12):
            if cap_idx - offset < 0:
                break
            prev_elem = paragraphs[cap_idx - offset]._element
            if has_drawing(prev_elem):
                new_p = doc.add_paragraph()
                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_p.paragraph_format.first_line_indent = Cm(0)
                run = new_p.add_run()
                run.add_picture(img_path, width=Cm(14))
                new_p_elem = new_p._element

                parent = prev_elem.getparent()
                parent.insert(parent.index(prev_elem), new_p_elem)
                parent.remove(prev_elem)

                replaced += 1
                print(f'  替换图片: {key} (段落[{cap_idx - offset}])')
                break

    last_p_elem = body[-1]
    if last_p_elem.tag == qn('w:p') and not last_p_elem.text:
        pass

    doc.save(OUT)
    print(f'\n完成！替换了 {replaced} 个代码截图')
    print(f'输出: {OUT}')


if __name__ == '__main__':
    main()
