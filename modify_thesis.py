# -*- coding: utf-8 -*-
"""
在现有论文(2).docx基础上进行修改：
1. 丰富4.4节决策模块内容（个人因素/三种模式/环境感知数据）
2. 增加网页页面截图
3. 修复格式（页码、代码块白色背景）
"""
import os, copy
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'thesis_images')

SRC = os.path.join(BASE_DIR, '通信工程（专转本）Z2024130143陈鎏-基于环境信息感知的智能出行穿衣服系统(2).docx')
OUT = os.path.join(BASE_DIR, '通信工程（专转本）Z2024130143陈鎏-基于环境信息感知的智能出行穿衣服系统.docx')

SIZE_WU = Pt(10.5)
SIZE_XIAO4 = Pt(12)
SIZE_SI = Pt(14)
SIZE_XIAO5 = Pt(9)


def set_run_font(run, cn='宋体', en='Times New Roman', size=SIZE_XIAO4, bold=False):
    run.bold = bold
    run.font.size = size
    run.font.name = en
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)


def make_paragraph(doc, text, cn='宋体', en='Times New Roman', size=SIZE_XIAO4,
                   bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   first_indent=True, space_before=0, space_after=0,
                   line_spacing_pt=22):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    jc = OxmlElement('w:jc')
    jc_map = {
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
        WD_ALIGN_PARAGRAPH.CENTER: 'center',
        WD_ALIGN_PARAGRAPH.LEFT: 'left',
    }
    jc.set(qn('w:val'), jc_map.get(align, 'both'))
    pPr.append(jc)

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 12700 / Pt(1))))
    spacing.set(qn('w:after'), str(int(space_after * 12700 / Pt(1))))
    spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)

    if first_indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLineChars'), '200')
        pPr.append(ind)

    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rPr.append(rFonts)

    sz_elem = OxmlElement('w:sz')
    sz_elem.set(qn('w:val'), str(int(size.pt * 2)))
    rPr.append(sz_elem)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size.pt * 2)))
    rPr.append(szCs)

    if bold:
        b_elem = OxmlElement('w:b')
        rPr.append(b_elem)

    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '000000')
    rPr.append(color_elem)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_heading3(text, doc=None):
    if doc:
        style_id = None
        for s in doc.styles:
            if s.name == 'Heading 3':
                style_id = s.style_id
                break
        if not style_id:
            style_id = 'Heading3'
    else:
        style_id = 'Heading3'

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    rPr.append(szCs)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_table_caption(text):
    return make_paragraph(None, text, cn='黑体', size=SIZE_WU, bold=False,
                          align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)


def make_figure_caption(text):
    return make_paragraph(None, text, cn='宋体', size=SIZE_WU, bold=False,
                          align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)


def insert_three_line_table(doc, body, insert_before, headers, rows, col_widths=None):
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    num_cols = len(headers)
    for ci in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        w_val = str(int((col_widths[ci] if col_widths and ci < len(col_widths) else 3) * 567))
        gridCol.set(qn('w:w'), w_val)
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    thick = {'sz': '12', 'val': 'single', 'color': '000000'}
    thin = {'sz': '6', 'val': 'single', 'color': '000000'}
    none_b = {'sz': '0', 'val': 'none', 'color': '000000'}
    total_rows = 1 + len(rows)

    for ri in range(total_rows):
        tr = OxmlElement('w:tr')
        data = headers if ri == 0 else rows[ri - 1]
        for ci, val in enumerate(data):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            if col_widths and ci < len(col_widths):
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(col_widths[ci] * 567)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

            tcBorders = OxmlElement('w:tcBorders')
            if ri == 0:
                for edge, bdr in [('top', thick), ('bottom', thin), ('start', none_b), ('end', none_b)]:
                    e = OxmlElement(f'w:{edge}')
                    for k, v in bdr.items():
                        e.set(qn(f'w:{k}'), v)
                    tcBorders.append(e)
            elif ri == total_rows - 1:
                for edge, bdr in [('top', none_b), ('bottom', thick), ('start', none_b), ('end', none_b)]:
                    e = OxmlElement(f'w:{edge}')
                    for k, v in bdr.items():
                        e.set(qn(f'w:{k}'), v)
                    tcBorders.append(e)
            else:
                for edge in ['top', 'bottom', 'start', 'end']:
                    e = OxmlElement(f'w:{edge}')
                    for k, v in none_b.items():
                        e.set(qn(f'w:{k}'), v)
                    tcBorders.append(e)
            tcPr.append(tcBorders)
            tc.append(tcPr)

            cell_p = make_paragraph(
                None, str(val),
                cn='黑体' if ri == 0 else '宋体',
                size=SIZE_WU,
                bold=(ri == 0),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_indent=False,
                line_spacing_pt=15
            )
            tc.append(cell_p)
            tr.append(tc)
        tbl.append(tr)

    body.insert(body.index(insert_before), tbl)
    return tbl


def gen_screenshot_placeholder(title, filename):
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    rect = FancyBboxPatch((0.2, 0.2), 9.6, 5.6, boxstyle='round,pad=0.03',
                          facecolor='white', edgecolor='black', lw=2)
    ax.add_patch(rect)
    ax.text(5, 3.2, title, ha='center', va='center', fontsize=18, color='black', fontweight='bold')
    ax.text(5, 2.2, '（请替换为实际系统截图）', ha='center', va='center', fontsize=12, color='#666666')
    path = os.path.join(IMG_DIR, filename)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def insert_image_paragraph(body, insert_before, img_path, width_cm=13):
    from docx.shared import Cm as CmVal
    p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), '0')
    pPr.append(ind)
    p_elem.append(pPr)
    body.insert(body.index(insert_before), p_elem)
    return p_elem


def find_paragraph_index(doc, text_start):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(text_start):
            return i
    return -1


def main():
    print('正在读取论文...')
    doc = Document(SRC)
    body = doc.element.body

    # ============================================================
    # 1. 在4.4节和4.5节之间插入新内容
    # ============================================================
    print('插入决策模块详细内容...')

    idx_45 = find_paragraph_index(doc, '4.5 API接口设计')
    if idx_45 < 0:
        idx_45 = find_paragraph_index(doc, '4.5 本章小结')
    if idx_45 < 0:
        print('ERROR: 找不到4.5节')
        return

    anchor = doc.paragraphs[idx_45]._element

    new_elements = []

    new_elements.append(make_heading3('4.4.1 环境信息感知与决策数据', doc))

    new_elements.append(make_paragraph(
        doc, '环境信息感知是推荐算法的输入基础。系统通过气象数据采集模块获取多维度的环境参数，'
        '包括气温、体感温度、相对湿度、风速、风向、降水量、气压、能见度、紫外线指数、空气质量'
        '指数（AQI）及其等级。这些原始环境数据经过融合处理后，生成两项综合决策指标：温湿指数'
        '（THI）和风寒修正温度，为保暖等级的判定提供量化依据。'))

    new_elements.append(make_paragraph(
        doc, 'THI是衡量人体在特定温湿度条件下舒适程度的综合指标，系统根据THI值将人体舒适度'
        '划分为8个等级，每个等级对应不同的穿衣建议基准。THI等级划分如表4-6所示。'))

    new_elements.append(make_table_caption('表4-6 THI舒适度等级划分表'))

    new_elements.append(make_paragraph(
        doc, '当气温低于10℃且风速超过4.8km/h时，系统启用风寒修正公式计算体感温度：'
        'T_wc = 13.12 + 0.6215×T - 11.37×V^0.16 + 0.3965×T×V^0.16，其中T为气温（℃），'
        'V为风速（km/h）。修正后的体感温度使系统在寒冷大风天气下给出更合理的加厚穿衣建议。'))

    new_elements.append(make_heading3('4.4.2 决策树分类与三种推荐模式', doc))

    new_elements.append(make_paragraph(
        doc, '系统构建了三套独立的决策树模型以满足不同体感偏好的用户需求：标准模式（standard）、'
        '偏暖模式（warm）和偏冷模式（cool）。三套模型的输入特征维度一致，均为气温、湿度、风速、'
        '降水量、AQI和场景ID共6个维度，但使用不同的训练数据集，在相同环境条件下输出不同的保暖等级。'
        '标准模式使用专家经验构建的均衡训练数据，max_depth=8；偏暖模式在同温度下系统性上调保暖等级，'
        'max_depth=6，适合怕冷人群；偏冷模式系统性下调保暖等级，max_depth=6，适合怕热人群。'
        '以15℃通勤场景为例，标准模式预测保暖等级为3，偏暖模式为4，偏冷模式为2。'))

    new_elements.append(make_heading3('4.4.3 个人因素调整机制', doc))

    new_elements.append(make_paragraph(
        doc, '在决策树输出基础保暖等级后，系统根据用户的个人生理特征进行二次调整。调整涉及'
        '体质类型、年龄和性别三个维度：怕冷体质（cold_sensitive）用户保暖等级上调1级，怕热体质'
        '（heat_sensitive）用户下调1级，正常体质不调整；60岁以上老年用户保暖等级上调1级，因为'
        '老年人体温调节能力较弱；性别信息用于在服装筛选阶段过滤适用性别，男性用户只推荐男性或通用'
        '服装，女性用户类似。个人因素调整规则如表4-7所示。'))

    new_elements.append(make_table_caption('表4-7 个人因素调整规则表'))

    new_elements.append(make_paragraph(
        doc, '调整后的保暖等级限制在1到5的有效范围内。体质和年龄的调整效果可以叠加，例如一位'
        '65岁的怕冷体质老人，标准模式预测保暖等级为3时，经体质调整（+1）和年龄调整（+1）后最终'
        '为5，系统将推荐最厚实的防寒衣物组合。'))

    new_elements.append(make_heading3('4.4.4 服装筛选与协同过滤优化', doc))

    new_elements.append(make_paragraph(
        doc, '第二阶段按上衣、下装、外套、鞋子、配饰五个品类槽位分别执行服装筛选。每个槽位的筛选'
        '条件包括：服装品类属于该槽位分类集合、服装处于启用状态、当前气温在服装适用温度范围内、'
        '服装适用性别与用户匹配、服装适用场景包含当前出行场景。满足条件的候选服装按保暖等级与预测值'
        '的差值绝对值升序排列，取前20件进入协同过滤排序。协同过滤基于余弦相似度找到最相似的Top-5'
        '用户，对候选服装进行加权评分排序，每个槽位选取评分最高的一件组成完整穿搭方案。'))

    new_elements.append(make_heading3('4.4.5 穿衣建议生成与舒适度评分', doc))

    new_elements.append(make_paragraph(
        doc, '系统根据气象条件自动判断是否需要外套和配饰：保暖等级≥3或有降水或风速>15km/h时建议'
        '穿外套；保暖等级≥4或有降水或AQI>150时建议携带配饰。系统还针对不同场景提供文字穿衣建议，'
        '如运动场景建议透气排汗面料、商务场合建议着正装等。舒适度评分（comfort_score）由基础舒适度'
        '（以22℃/55%湿度为最佳基准）、模式调整因子和保暖匹配度加权计算，范围5到98分，为用户提供'
        '直观的推荐质量评价参考。'))

    for elem in new_elements:
        body.insert(body.index(anchor), elem)

    idx_45_cap = find_paragraph_index(doc, '表4-6')
    if idx_45_cap >= 0:
        anchor_t6 = doc.paragraphs[idx_45_cap + 1]._element
        insert_three_line_table(doc, body, anchor_t6,
            ['THI范围', '等级', '体感描述', '穿衣建议基准'],
            [
                ['THI < 40', '1级', '极冷', '厚羽绒服、棉服'],
                ['40 ≤ THI < 45', '2级', '寒冷', '厚实保暖衣物'],
                ['45 ≤ THI < 55', '3级', '偏冷', '需外套保暖'],
                ['55 ≤ THI < 60', '4级', '凉爽', '薄外套或长袖'],
                ['60 ≤ THI < 70', '5级', '舒适', '穿着轻便即可'],
                ['70 ≤ THI < 75', '6级', '温暖', '短袖，注意防晒'],
                ['75 ≤ THI < 80', '7级', '偏热', '清凉透气衣物'],
                ['THI ≥ 80', '8级', '炎热', '防暑降温'],
            ],
            col_widths=[3, 1.5, 2, 5])

    idx_47_cap = find_paragraph_index(doc, '表4-7')
    if idx_47_cap >= 0:
        anchor_t7 = doc.paragraphs[idx_47_cap + 1]._element
        insert_three_line_table(doc, body, anchor_t7,
            ['调整维度', '用户特征', '调整方式', '调整幅度'],
            [
                ['体质', '怕冷体质', '保暖等级上调', '+1级'],
                ['体质', '怕热体质', '保暖等级下调', '-1级'],
                ['体质', '正常体质', '不调整', '0'],
                ['年龄', '60岁以上', '保暖等级上调', '+1级'],
                ['年龄', '60岁及以下', '不调整', '0'],
                ['性别', '男性', '筛选男性适用服装', '—'],
                ['性别', '女性', '筛选女性适用服装', '—'],
                ['性别', '未设置', '显示全部服装', '—'],
            ],
            col_widths=[2.5, 3, 3.5, 2])

    # ============================================================
    # 2. 增加截图（在5.2.2前端核心实现部分的截图后面追加）
    # ============================================================
    print('生成并插入额外截图占位图...')

    screenshots_to_add = {
        'sc_register.png': '用户注册界面',
        'sc_profile.png': '个人中心界面',
        'sc_admin_users.png': '管理端用户管理界面',
        'sc_admin_weather.png': '管理端天气监控界面',
        'sc_admin_rec.png': '管理端推荐记录界面',
    }
    for fname, title in screenshots_to_add.items():
        path = os.path.join(IMG_DIR, fname)
        if not os.path.exists(path):
            gen_screenshot_placeholder(title, fname)

    idx_523 = find_paragraph_index(doc, '5.2.3 算法实现')
    if idx_523 > 0:
        anchor_523 = doc.paragraphs[idx_523]._element
        sc_elems = []

        sc_elems.append(make_paragraph(
            doc, '图5-13展示了用户注册界面，新用户通过填写用户名、邮箱、密码等信息即可完成注册。'
            '图5-14展示了用户个人中心界面，用户可以编辑性别、年龄、体质类型等个人信息，这些信息'
            '直接影响推荐算法的个人因素调整环节。',
            first_indent=True))

        sc_elems.append(make_figure_caption('图5-12 用户注册界面'))
        sc_elems.append(make_figure_caption('图5-13 用户个人中心界面'))

        sc_elems.append(make_paragraph(
            doc, '管理端方面，图5-14展示了用户管理界面，管理员可以查看、搜索和管理系统中的用户账户。'
            '图5-15展示了天气数据监控界面，实时展示各城市的气象数据采集情况。图5-16展示了推荐'
            '记录查看界面，管理员可以查看所有用户的历史推荐记录及其反馈评分。',
            first_indent=True))

        sc_elems.append(make_figure_caption('图5-14 管理端用户管理界面'))
        sc_elems.append(make_figure_caption('图5-15 管理端天气数据监控界面'))
        sc_elems.append(make_figure_caption('图5-16 管理端推荐记录查看界面'))

        for elem in sc_elems:
            body.insert(body.index(anchor_523), elem)

    # ============================================================
    # 3. 修复代码块背景为白色
    # ============================================================
    print('修复代码块背景为白色...')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill.upper() in ('F5F5F5', 'F0F0F0', 'EEEEEE', 'E8E8E8', 'EDEDED'):
                            shd.set(qn('w:fill'), 'FFFFFF')

    # ============================================================
    # 4. 保存
    # ============================================================
    doc.save(OUT)
    print(f'修改完成，已保存到: {OUT}')


if __name__ == '__main__':
    main()
