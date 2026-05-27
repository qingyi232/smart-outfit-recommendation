# -*- coding: utf-8 -*-
"""
直接操作原始 docx 文件修复所有批注问题
"""
import zipfile
import shutil
import os
import re
from xml.etree import ElementTree as ET

src_path = r'F:\26毕设2\基于环境信息感知的智能出行穿衣服务系统\通信工程（专转本）Z2024130143陈鎏-毕业论文-.docx'
dst_path = r'F:\26毕设2\基于环境信息感知的智能出行穿衣服务系统\通信工程（专转本）Z2024130143陈鎏-毕业论文-修改版.docx'

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

ET.register_namespace('w', NS_W)
ET.register_namespace('r', NS_R)

shutil.copy2(src_path, dst_path)

from docx import Document
doc = Document(dst_path)

changes = {'ref': 0, 'fig_title': 0, 'formula': 0}
chapter = 0

formula_nums = {
    'THI = (1.8': '2-1',
    'sim(u, v) = cos': '2-2',
    'T(d) = T_base': '5-1',
}

for para in doc.paragraphs:
    full_text = para.text.strip()
    
    ch_match = re.match(r'^(\d+)\s+\S', full_text)
    if ch_match and len(full_text) < 40:
        chapter = int(ch_match.group(1))
    
    # 批注0: 参考文献上标
    for run in para.runs:
        if run.text and re.search(r'\[\d+\]', run.text):
            run.font.superscript = True
            changes['ref'] += 1
    
    # 批注3: 图名格式统一
    if re.match(r'^图\s*\d', full_text) and len(full_text) < 80:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        from docx.oxml.ns import qn
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        changes['fig_title'] += 1
    
    # 批注1: 公式编号
    for key, num in formula_nums.items():
        if key in full_text and len(full_text) < 120:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run('\t(' + num + ')')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            
            pPr = para._element.get_or_add_pPr()
            tabs = pPr.makeelement(qn('w:tabs'), {})
            tab = tabs.makeelement(qn('w:tab'), {
                qn('w:val'): 'right',
                qn('w:pos'): '8306'
            })
            tabs.append(tab)
            pPr.append(tabs)
            changes['formula'] += 1

doc.save(dst_path)

# 删除批注
import zipfile as zf
import tempfile

temp_path = dst_path + '.tmp'
with zf.ZipFile(dst_path, 'r') as zin:
    with zf.ZipFile(temp_path, 'w') as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/comments.xml':
                root = ET.fromstring(data)
                for comment in root.findall('.//{%s}comment' % NS_W):
                    root.remove(comment)
                data = ET.tostring(root, encoding='unicode').encode('utf-8')
                print('已清除所有批注')
            
            if item.filename == 'word/document.xml':
                content = data.decode('utf-8')
                # 删除批注引用标记
                content = re.sub(r'<w:commentRangeStart[^/]*/>', '', content)
                content = re.sub(r'<w:commentRangeEnd[^/]*/>', '', content)
                content = re.sub(r'<w:commentReference[^/]*/>', '', content)
                # 删除批注run
                content = re.sub(r'<w:r>\s*<w:rPr>.*?</w:rPr>\s*<w:commentReference[^/]*/>\s*</w:r>', '', content, flags=re.DOTALL)
                data = content.encode('utf-8')
                print('已清除文档中的批注标记')
            
            zout.writestr(item, data)

os.replace(temp_path, dst_path)

print(f'\n修改汇总:')
print(f'  参考文献上标: {changes["ref"]} 处')
print(f'  图名格式统一: {changes["fig_title"]} 处')
print(f'  公式编号添加: {changes["formula"]} 处')
print(f'  批注已全部清除')
print(f'\n保存到: {dst_path}')
