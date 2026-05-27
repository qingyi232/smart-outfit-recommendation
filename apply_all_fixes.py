# -*- coding: utf-8 -*-
"""在已修复上标的版本上应用其他所有修改"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, zipfile, os, shutil
from xml.etree import ElementTree as ET

docx_path = r'F:\26毕设2\基于环境信息感知的智能出行穿衣服务系统\通信工程（专转本）Z2024130143陈鎏-毕业论文-修改版.docx'

doc = Document(docx_path)

# ============ 批注3: 图名格式统一 ============
fig_count = 0
for para in doc.paragraphs:
    full_text = para.text.strip()
    if re.match(r'^图\s*\d', full_text) and len(full_text) < 80:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        fig_count += 1

# ============ 批注1: 公式编号 ============
formula_nums = {
    'THI = (1.8': '2-1',
    'sim(u, v) = cos': '2-2',
    'T(d) = T_base': '5-1',
}
formula_count = 0
for para in doc.paragraphs:
    full_text = para.text.strip()
    for key, num in formula_nums.items():
        if key in full_text and len(full_text) < 120 and '(' + num + ')' not in full_text:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run('\t(' + num + ')')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            pPr = para._element.get_or_add_pPr()
            tabs = pPr.makeelement(qn('w:tabs'), {})
            tab = tabs.makeelement(qn('w:tab'), {qn('w:val'): 'right', qn('w:pos'): '8306'})
            tabs.append(tab)
            pPr.append(tabs)
            formula_count += 1

# ============ 批注2: 插入本章小结 ============
summaries = {
    2: ('2.4 本章小结', '在本章的论述当中，主要是对与本系统的设计和开发工作有着比较密切关联的几项关键性技术进行了相对较为系统的梳理和介绍。首先在环境数据采集这一块，对气象数据的采集技术手段以及多种不同来源的数据进行融合处理的基本方法做了一些说明和讨论。在这之后，本章又对本系统所用到的两个比较核心的推荐算法进行了原理层面上的介绍和分析，一个是用来进行保暖等级预测判断的决策树分类算法，另一个是用来实现面向不同用户的个性化推荐优化功能的协同过滤算法。最后，在开发框架技术方面，本章分别对后端开发所采用的Python Flask框架以及前端开发所使用的Vue.js框架的技术特点和适用性做了一定程度的阐述。通过本章对上述相关技术的综述和分析，可以为后面章节中将要展开的系统需求分析、架构设计以及功能实现等工作提供必要的理论知识基础和技术方面的支撑依据。'),
    3: ('3.4 本章小结', '在本章节的内容里面，本文主要是从系统使用者的角度出发，对智能出行穿衣服务系统在功能方面和非功能方面所需要满足的各项需求进行了比较详细的调查研究和分析整理工作。在用户需求分析的部分，本文将系统的使用者划分成了普通用户和管理员这两种不同的角色类型，并且分别针对每一种角色各自的使用场景和操作需要进行了需求的归纳和总结，同时还绘制了对应的用例图来对两类用户的主要交互行为做出了直观的展示。在功能需求方面，本文明确了包括天气信息查询、智能穿衣推荐、用户偏好管理、推荐历史查看以及后台管理等在内的若干核心功能模块。在非功能需求方面，则从系统的性能响应速度、界面交互的易用性、数据的安全保护等几个维度提出了相应的设计要求和约束条件。这些需求分析的成果将作为下一章系统设计环节的重要输入依据。'),
    4: ('4.5 本章小结', '本章作为整篇论文中在技术方案设计方面最为核心的一个章节，围绕着智能出行穿衣服务系统的各个组成部分展开了比较详尽和深入的设计工作。在系统的整体架构设计上面，本文选择采用了前后端分离模式下的B/S架构方案，后端服务基于Flask框架来构建RESTful风格的API接口，前端界面则使用Vue 3框架进行开发实现。在数据库设计环节，先后完成了概念结构层面的E-R图的绘制以及逻辑结构层面的数据表设计工作，确定了用户表、服装表、天气记录表和推荐记录表等几张核心的数据表及其字段结构。在功能模块设计方面，本文对系统的各个子功能模块进行了划分和详细设计。在推荐算法设计部分，设计了基于决策树进行保暖等级预测并结合协同过滤进行个性化优化的推荐算法流程。上述设计方案为后续章节中系统的编码实现工作提供了清晰完整的技术蓝图。'),
    5: ('5.4 本章小结', '在本章的论述过程当中，本文对智能出行穿衣服务系统的编码实现过程以及测试验证工作进行了较为详细的展示和分析说明。在系统实现部分，首先介绍了项目的开发环境配置情况，然后分别从后端核心实现、前端核心实现以及算法实现这三个方面对系统的关键代码逻辑和实现细节做了阐述，并且配合代码截图和界面截图来帮助读者更好地理解各个功能模块的具体实现方式。在测试部分，本文设计了包含多个测试用例的功能测试方案，对系统的用户注册登录、天气数据获取、穿衣推荐生成、服装管理操作等核心功能逐一进行了测试验证，同时还开展了针对系统响应速度和并发处理能力的性能测试工作。从测试结果来看，系统的各项主要功能均能够按照预期正常地运行工作，在性能指标方面也基本达到了设计阶段所提出的要求标准。'),
    6: ('6.2 本章小结', '在本章的研究工作中，本文主要是对已经开发完成的智能出行穿衣服务系统开展了一次较为全面的用户体验评估活动。为了能够收集到比较真实和有参考价值的评估数据，本文设计了专门的用户体验评估方案，并且邀请了若干名测试参与人员对系统进行了实际的操作使用和体验。在评估过程中，参与测试的人员分别从系统的界面美观程度、操作流畅性、推荐结果准确性、功能完整性以及整体满意度等多个不同的维度给出了各自的主观评分和文字反馈意见。从最终汇总得到的评估结果数据来看，大多数的参与者对于系统的整体表现给予了较为正面的评价，特别是在界面设计的友好性和穿衣推荐结果的参考价值方面获得了比较高的认可度，但同时也有部分测试者提出了一些关于推荐精准度和功能丰富度方面的改进意见和建议，这些反馈对于后续进一步完善和优化系统具有重要的参考价值。'),
}

chapter = 0
insert_positions = {}
for i, para in enumerate(doc.paragraphs):
    full_text = para.text.strip()
    ch_match = re.match(r'^(\d+)\s+\S', full_text)
    if ch_match and len(full_text) < 40:
        new_ch = int(ch_match.group(1))
        if chapter in summaries and chapter != new_ch:
            insert_positions[chapter] = i - 1
        chapter = new_ch

if 6 in summaries and 6 not in insert_positions:
    for i, para in enumerate(doc.paragraphs):
        if re.match(r'^7\s', para.text.strip()):
            insert_positions[6] = i - 1
            break

summary_count = 0
for ch in sorted(insert_positions.keys(), reverse=True):
    pos = insert_positions[ch]
    title, text = summaries[ch]
    ref_para = doc.paragraphs[pos]
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = '黑体'
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    body_para = doc.add_paragraph()
    body_run = body_para.add_run('\u3000\u3000' + text)
    body_run.font.name = '宋体'
    body_run.font.size = Pt(12)
    body_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    ref_element = ref_para._element
    ref_element.addnext(body_para._element)
    ref_element.addnext(title_para._element)
    summary_count += 1

doc.save(docx_path)

# 清除批注
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
temp_path = docx_path + '.tmp'
with zipfile.ZipFile(docx_path, 'r') as zin:
    with zipfile.ZipFile(temp_path, 'w') as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/comments.xml':
                root = ET.fromstring(data)
                for c in root.findall('.//{%s}comment' % NS_W):
                    root.remove(c)
                data = ET.tostring(root, encoding='unicode').encode('utf-8')
            if item.filename == 'word/document.xml':
                content = data.decode('utf-8')
                content = re.sub(r'<w:commentRangeStart[^/]*/>', '', content)
                content = re.sub(r'<w:commentRangeEnd[^/]*/>', '', content)
                content = re.sub(r'<w:commentReference[^/]*/>', '', content)
                data = content.encode('utf-8')
            zout.writestr(item, data)
os.replace(temp_path, docx_path)

print(f'图名格式: {fig_count} 处')
print(f'公式编号: {formula_count} 处')
print(f'本章小结: {summary_count} 章')
print(f'批注已清除')
print(f'保存到: {docx_path}')
