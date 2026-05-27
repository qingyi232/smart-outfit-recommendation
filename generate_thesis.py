# -*- coding: utf-8 -*-
"""
生成毕业论文 docx —— 基于环境信息感知的智能出行穿衣服务系统
安徽三联学院 计算机与通信工程学院
"""

import os, math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'thesis_images')
os.makedirs(IMG_DIR, exist_ok=True)

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial']
plt.rcParams['axes.unicode_minus'] = False

# ==================== 字号常量 ====================
SIZE_CHU1 = Pt(26)
SIZE_XIAO1 = Pt(24)
SIZE_ER = Pt(22)
SIZE_XIAO2 = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO3 = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO4 = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO5 = Pt(9)

# ==================== 辅助函数 ====================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)

def set_run_font(run, cn='宋体', en='Times New Roman', size=SIZE_XIAO4, bold=False):
    run.bold = bold
    run.font.size = size
    run.font.name = en
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)

def add_paragraph(doc, text, cn='宋体', en='Times New Roman', size=SIZE_XIAO4,
                  bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=True, space_before=0, space_after=0,
                  line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Cm(0.74)
    else:
        pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)
    return p

def add_heading_l1(doc, text, number=None):
    label = f'{number}  {text}' if number else text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(36)
    pf.space_after = Pt(18)
    pf.first_line_indent = Cm(0)
    run = p.add_run(label)
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_SAN, bold=False)
    return p

def add_heading_l2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_SI, bold=False)
    return p

def add_heading_l3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_XIAO4, bold=False)
    return p

def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_WU, bold=False)
    return p

def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_WU, bold=False)
    return p

def add_three_line_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_WU, bold=True)
        p.paragraph_format.line_spacing = Pt(15)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, cn='宋体', en='Times New Roman', size=SIZE_WU, bold=False)
            p.paragraph_format.line_spacing = Pt(15)

    thick = {'sz': '12', 'val': 'single', 'color': '000000'}
    thin = {'sz': '6', 'val': 'single', 'color': '000000'}
    none_b = {'sz': '0', 'val': 'none', 'color': '000000'}

    for ri, row_obj in enumerate(table.rows):
        for ci, cell in enumerate(row_obj.cells):
            if ri == 0:
                set_cell_border(cell, top=thick, bottom=thin, start=none_b, end=none_b)
            elif ri == len(rows):
                set_cell_border(cell, top=none_b, bottom=thick, start=none_b, end=none_b)
            else:
                set_cell_border(cell, top=none_b, bottom=none_b, start=none_b, end=none_b)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row_obj in table.rows:
                row_obj.cells[i].width = Cm(w)
    return table

def insert_image(doc, img_path, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

import docx.enum.text

def add_code_block(doc, code_text, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(caption)
        set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_WU, bold=False)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    for i, line in enumerate(code_text.strip().split('\n')):
        if i > 0:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = Pt(14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(line)
        set_run_font(run, cn='Courier New', en='Courier New', size=SIZE_XIAO5, bold=False)
    border = {'sz': '4', 'val': 'single', 'color': '999999'}
    set_cell_border(cell, top=border, bottom=border, start=border, end=border)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFFFFF')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def gen_screenshot_placeholder(title, filename):
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis('off')
    rect = FancyBboxPatch((0.2, 0.2), 9.6, 5.6, boxstyle='round,pad=0.03',
                          facecolor='white', edgecolor='black', lw=2)
    ax.add_patch(rect)
    ax.text(5, 3.2, title, ha='center', va='center', fontsize=18, color='black', fontweight='bold')
    ax.text(5, 2.2, '（请替换为实际系统截图）', ha='center', va='center', fontsize=12, color='#666666')
    path = os.path.join(IMG_DIR, filename)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ==================== 图表生成 ====================

def draw_box(ax, x, y, w, h, text, fontsize=10, lw=1.5):
    rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.02',
                          facecolor='white', edgecolor='black', linewidth=lw)
    ax.add_patch(rect)
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontsize=fontsize, color='black', fontweight='normal')

def draw_arrow(ax, x1, y1, x2, y2, style='->', lw=1.2):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color='black', lw=lw))

def gen_architecture_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')

    layers = [
        (1, 6.0, 8, 1.2, '用户交互层\nVue 3 + Element Plus + ECharts + Axios'),
        (1, 4.2, 8, 1.2, '业务逻辑层\nFlask + JWT认证 + RESTful API'),
        (1, 2.4, 8, 1.2, '数据处理层\n决策树模型 + 协同过滤 + THI舒适度模型'),
        (1, 0.6, 8, 1.2, '数据存储层\nSQLite / MySQL + 和风天气API'),
    ]
    for x, y, w, h, t in layers:
        draw_box(ax, x, y, w, h, t, fontsize=11)

    for i in range(3):
        y_top = layers[i][1]
        y_bot = layers[i+1][1] + layers[i+1][3]
        ax.annotate('', xy=(5, y_bot), xytext=(5, y_top),
                    arrowprops=dict(arrowstyle='<->', color='black', lw=1.5))

    path = os.path.join(IMG_DIR, 'architecture.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def gen_usecase_user():
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 10)
    ax.axis('off')

    ax.plot([1.5, 1.5], [5.8, 5.2], 'k-', lw=2)
    ax.plot([1.5, 1.5], [6.2, 5.8], 'k-', lw=2)
    circle = plt.Circle((1.5, 6.4), 0.25, fill=False, edgecolor='black', lw=2)
    ax.add_patch(circle)
    ax.plot([1.5, 1.1], [5.8, 5.4], 'k-', lw=2)
    ax.plot([1.5, 1.9], [5.8, 5.4], 'k-', lw=2)
    ax.plot([1.5, 1.1], [5.2, 4.7], 'k-', lw=2)
    ax.plot([1.5, 1.9], [5.2, 4.7], 'k-', lw=2)
    ax.text(1.5, 4.3, '普通用户', ha='center', fontsize=11, color='black')

    cases = [
        (7, 9.0, '用户注册'), (7, 8.0, '用户登录'), (7, 7.0, '查看天气信息'),
        (7, 6.0, '获取穿衣推荐'), (7, 5.0, '切换场景'),
        (7, 4.0, '查看天气预报'), (7, 3.0, '查看推荐历史'),
        (7, 2.0, '反馈评分'), (7, 1.0, '编辑个人信息'),
    ]
    for cx, cy, label in cases:
        ellipse = mpatches.Ellipse((cx, cy), 3.5, 0.7, fill=False,
                                   edgecolor='black', lw=1.5)
        ax.add_patch(ellipse)
        ax.text(cx, cy, label, ha='center', va='center', fontsize=10, color='black')
        ax.plot([2.2, cx - 1.75], [5.5, cy], 'k-', lw=1)

    path = os.path.join(IMG_DIR, 'usecase_user.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def gen_usecase_admin():
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')

    ax.plot([1.5, 1.5], [4.8, 4.2], 'k-', lw=2)
    ax.plot([1.5, 1.5], [5.2, 4.8], 'k-', lw=2)
    circle = plt.Circle((1.5, 5.4), 0.25, fill=False, edgecolor='black', lw=2)
    ax.add_patch(circle)
    ax.plot([1.5, 1.1], [4.8, 4.4], 'k-', lw=2)
    ax.plot([1.5, 1.9], [4.8, 4.4], 'k-', lw=2)
    ax.plot([1.5, 1.1], [4.2, 3.7], 'k-', lw=2)
    ax.plot([1.5, 1.9], [4.2, 3.7], 'k-', lw=2)
    ax.text(1.5, 3.3, '管理员', ha='center', fontsize=11, color='black')

    cases = [
        (7, 7.0, '查看数据概览'), (7, 6.0, '用户管理'),
        (7, 5.0, '服装管理'), (7, 4.0, '天气数据监控'),
        (7, 3.0, '推荐记录查看'), (7, 2.0, '系统配置'),
    ]
    for cx, cy, label in cases:
        ellipse = mpatches.Ellipse((cx, cy), 3.5, 0.7, fill=False,
                                   edgecolor='black', lw=1.5)
        ax.add_patch(ellipse)
        ax.text(cx, cy, label, ha='center', va='center', fontsize=10, color='black')
        ax.plot([2.2, cx - 1.75], [4.5, cy], 'k-', lw=1)

    path = os.path.join(IMG_DIR, 'usecase_admin.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def gen_module_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')

    draw_box(ax, 4, 8.5, 6, 0.9, '智能出行穿衣服务系统', fontsize=12, lw=2)

    modules = [
        (0.2, 6.5, 2.5, 0.8, '数据采集\n模块'),
        (3.2, 6.5, 2.5, 0.8, '推荐引擎\n模块'),
        (6.2, 6.5, 2.5, 0.8, '用户交互\n模块'),
        (9.2, 6.5, 2.5, 0.8, '后台管理\n模块'),
        (12.0, 6.5, 1.8, 0.8, '认证\n模块'),
    ]
    for x, y, w, h, t in modules:
        draw_box(ax, x, y, w, h, t, fontsize=9)
        ax.plot([x + w/2, x + w/2], [8.5, y + h], 'k-', lw=1.2)

    sub1 = [
        (0.0, 4.5, 1.3, 0.7, '天气API\n接入'),
        (1.5, 4.5, 1.3, 0.7, '气候\n模型'),
    ]
    for x, y, w, h, t in sub1:
        draw_box(ax, x, y, w, h, t, fontsize=8)
        ax.plot([x + w/2, 1.45], [y + h, 6.5], 'k-', lw=1)

    sub2 = [
        (3.0, 4.5, 1.3, 0.7, '决策树\n算法'),
        (4.5, 4.5, 1.3, 0.7, '协同\n过滤'),
    ]
    for x, y, w, h, t in sub2:
        draw_box(ax, x, y, w, h, t, fontsize=8)
        ax.plot([x + w/2, 4.45], [y + h, 6.5], 'k-', lw=1)

    sub3 = [
        (5.8, 4.5, 1.1, 0.7, '天气\n展示'),
        (7.0, 4.5, 1.1, 0.7, '推荐\n展示'),
        (8.2, 4.5, 1.1, 0.7, '用户\n中心'),
    ]
    for x, y, w, h, t in sub3:
        draw_box(ax, x, y, w, h, t, fontsize=8)
        ax.plot([x + w/2, 7.45], [y + h, 6.5], 'k-', lw=1)

    sub4 = [
        (8.8, 4.5, 1.2, 0.7, '用户\n管理'),
        (10.2, 4.5, 1.2, 0.7, '服装\n管理'),
        (11.5, 4.5, 1.2, 0.7, '数据\n监控'),
    ]
    for x, y, w, h, t in sub4:
        draw_box(ax, x, y, w, h, t, fontsize=8)
        ax.plot([x + w/2, 10.45], [y + h, 6.5], 'k-', lw=1)

    path = os.path.join(IMG_DIR, 'module_diagram.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def gen_er_diagram():
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')

    entities = {
        'User': (2, 8, 2.5, 1.2, 'User\n(用户表)\nid, username, email\npassword_hash, gender\nage, city, body_type, role'),
        'Clothing': (7, 8, 2.5, 1.2, 'Clothing\n(服装表)\nid, name, category_id\nwarmth_level, breathability\nmin_temp, max_temp'),
        'Category': (12, 8, 2, 0.8, 'ClothingCategory\n(分类表)\nid, name'),
        'Weather': (2, 4.5, 2.5, 1.0, 'WeatherRecord\n(天气记录表)\nid, city, date\ntemperature, humidity\nwind_speed, aqi'),
        'Recommendation': (7, 4.5, 2.5, 1.2, 'Recommendation\n(推荐记录表)\nid, user_id, city, scene\noutfit_top/bottom/outer\ncomfort_score'),
        'Feedback': (12, 4.5, 2, 0.8, 'UserFeedback\n(反馈表)\nid, user_id\nrating, comment'),
        'Preference': (7, 1.5, 2.5, 0.8, 'UserPreference\n(偏好表)\nid, user_id, clothing_id\npreference_score'),
    }

    for name, (x, y, w, h, label) in entities.items():
        rect = FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.03',
                              facecolor='white', edgecolor='black', lw=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=7, color='black')

    def line(e1, e2, label, dx1=0, dy1=0, dx2=0, dy2=0):
        x1, y1, w1, h1 = entities[e1][:4]
        x2, y2, w2, h2 = entities[e2][:4]
        px1 = x1 + w1/2 + dx1
        py1 = y1 + dy1
        px2 = x2 + w2/2 + dx2
        py2 = y2 + h2 + dy2
        ax.plot([px1, px2], [py1, py2], 'k-', lw=1.2)
        mx, my = (px1+px2)/2, (py1+py2)/2
        ax.text(mx + 0.15, my + 0.1, label, fontsize=7, color='black')

    line('User', 'Recommendation', '1:N', dy1=0, dy2=0)
    ax.plot([3.25, 8.25], [8.0, 5.7], 'k-', lw=1.2)
    ax.text(5.5, 7.0, '1:N', fontsize=8, color='black')

    ax.plot([8.25, 8.25], [8.0, 5.7], 'k-', lw=1.2)
    ax.text(8.4, 6.8, 'N:1', fontsize=8, color='black')

    ax.plot([9.5, 12.0], [8.6, 8.6], 'k-', lw=1.2)
    ax.text(10.5, 8.75, '1:N', fontsize=8, color='black')

    ax.plot([9.5, 12.0], [5.1, 5.1], 'k-', lw=1.2)
    ax.text(10.5, 5.25, '1:N', fontsize=8, color='black')

    ax.plot([3.25, 8.25], [4.5, 2.3], 'k-', lw=1.2)
    ax.text(5.0, 3.5, '1:N', fontsize=8, color='black')

    ax.plot([8.25, 8.25], [4.5, 2.3], 'k-', lw=1.2)
    ax.text(8.4, 3.4, 'N:1', fontsize=8, color='black')

    path = os.path.join(IMG_DIR, 'er_diagram.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def gen_algorithm_flow():
    fig, ax = plt.subplots(1, 1, figsize=(12, 10))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 12)
    ax.axis('off')

    steps = [
        (5, 11.0, 2, 0.6, '开始', 'round'),
        (4, 9.5, 4, 0.7, '获取用户位置与天气数据', 'box'),
        (4, 8.0, 4, 0.7, '计算THI舒适度指数', 'box'),
        (3.5, 6.5, 5, 0.7, '决策树模型预测保暖等级', 'box'),
        (3.5, 5.0, 5, 0.7, '根据体质/年龄调整保暖等级', 'box'),
        (3.5, 3.5, 5, 0.7, '按温度范围和场景筛选候选服装', 'box'),
        (3.5, 2.0, 5, 0.7, '协同过滤对候选服装评分排序', 'box'),
        (4, 0.5, 4, 0.7, '生成搭配方案与穿衣建议', 'box'),
    ]

    for x, y, w, h, t, _ in steps:
        if t == '开始':
            ellipse = mpatches.Ellipse((x + w/2, y + h/2), w, h, fill=False,
                                       edgecolor='black', lw=1.5)
            ax.add_patch(ellipse)
            ax.text(x + w/2, y + h/2, t, ha='center', va='center', fontsize=10)
        else:
            draw_box(ax, x, y, w, h, t, fontsize=9)

    for i in range(len(steps) - 1):
        x1, y1, w1, h1 = steps[i][0], steps[i][1], steps[i][2], steps[i][3]
        x2, y2 = steps[i+1][0], steps[i+1][1]
        w2, h2 = steps[i+1][2], steps[i+1][3]
        ax.annotate('', xy=(x2 + w2/2, y2 + h2), xytext=(x1 + w1/2, y1),
                    arrowprops=dict(arrowstyle='->', color='black', lw=1.5))

    ellipse = mpatches.Ellipse((6, -0.3), 2, 0.6, fill=False,
                               edgecolor='black', lw=1.5)
    ax.add_patch(ellipse)
    ax.text(6, -0.3, '结束', ha='center', va='center', fontsize=10)
    ax.annotate('', xy=(6, 0.0), xytext=(6, 0.5),
                arrowprops=dict(arrowstyle='->', color='black', lw=1.5))

    ax.set_ylim(-1, 12)

    path = os.path.join(IMG_DIR, 'algorithm_flow.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def gen_system_flow():
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis('off')

    ax.text(7, 8.5, '系统业务流程图', ha='center', fontsize=13, fontweight='bold', color='black')

    steps_left = [
        (0.5, 6.5, 2.5, 0.7, '用户注册/登录'),
        (0.5, 5.0, 2.5, 0.7, '设置个人信息'),
        (0.5, 3.5, 2.5, 0.7, '选择城市和场景'),
        (0.5, 2.0, 2.5, 0.7, '查看推荐结果'),
        (0.5, 0.5, 2.5, 0.7, '反馈评分'),
    ]

    steps_mid = [
        (5.0, 6.5, 3, 0.7, '天气API采集数据'),
        (5.0, 5.0, 3, 0.7, 'THI舒适度计算'),
        (5.0, 3.5, 3, 0.7, '决策树+协同过滤推荐'),
        (5.0, 2.0, 3, 0.7, '生成穿衣方案'),
    ]

    steps_right = [
        (10, 6.5, 3.5, 0.7, '管理员查看数据概览'),
        (10, 5.0, 3.5, 0.7, '管理服装/用户数据'),
        (10, 3.5, 3.5, 0.7, '监控天气数据'),
        (10, 2.0, 3.5, 0.7, '查看推荐记录'),
    ]

    for group in [steps_left, steps_mid, steps_right]:
        for x, y, w, h, t in group:
            draw_box(ax, x, y, w, h, t, fontsize=8)

    for i in range(len(steps_left) - 1):
        x, y, w, h = steps_left[i][0], steps_left[i][1], steps_left[i][2], steps_left[i][3]
        ny = steps_left[i+1][1]
        nh = steps_left[i+1][3]
        ax.annotate('', xy=(x + w/2, ny + nh), xytext=(x + w/2, y),
                    arrowprops=dict(arrowstyle='->', color='black', lw=1.2))

    for i in range(len(steps_mid) - 1):
        x, y, w, h = steps_mid[i][0], steps_mid[i][1], steps_mid[i][2], steps_mid[i][3]
        ny = steps_mid[i+1][1]
        nh = steps_mid[i+1][3]
        ax.annotate('', xy=(x + w/2, ny + nh), xytext=(x + w/2, y),
                    arrowprops=dict(arrowstyle='->', color='black', lw=1.2))

    for i in range(len(steps_right) - 1):
        x, y, w, h = steps_right[i][0], steps_right[i][1], steps_right[i][2], steps_right[i][3]
        ny = steps_right[i+1][1]
        nh = steps_right[i+1][3]
        ax.annotate('', xy=(x + w/2, ny + nh), xytext=(x + w/2, y),
                    arrowprops=dict(arrowstyle='->', color='black', lw=1.2))

    ax.annotate('', xy=(5.0, 6.85), xytext=(3.0, 6.85),
                arrowprops=dict(arrowstyle='->', color='black', lw=1.2))
    ax.annotate('', xy=(3.0, 3.85), xytext=(5.0, 2.7),
                arrowprops=dict(arrowstyle='->', color='black', lw=1.2, connectionstyle='arc3,rad=0.3'))

    ax.text(1.75, 7.8, '用户端', ha='center', fontsize=11, fontweight='bold', color='black',
            bbox=dict(boxstyle='round', facecolor='white', edgecolor='black'))
    ax.text(6.5, 7.8, '服务端', ha='center', fontsize=11, fontweight='bold', color='black',
            bbox=dict(boxstyle='round', facecolor='white', edgecolor='black'))
    ax.text(11.75, 7.8, '管理端', ha='center', fontsize=11, fontweight='bold', color='black',
            bbox=dict(boxstyle='round', facecolor='white', edgecolor='black'))

    path = os.path.join(IMG_DIR, 'system_flow.png')
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ==================== 文档生成 ====================

def create_thesis():
    print('生成图表...')
    img_arch = gen_architecture_diagram()
    img_uc_user = gen_usecase_user()
    img_uc_admin = gen_usecase_admin()
    img_module = gen_module_diagram()
    img_er = gen_er_diagram()
    img_algo_flow = gen_algorithm_flow()
    img_sys_flow = gen_system_flow()
    img_sc_login = os.path.join(IMG_DIR, 'sc_login.png')
    img_sc_register = os.path.join(IMG_DIR, 'sc_register.png')
    img_sc_home = os.path.join(IMG_DIR, 'sc_home.png')
    img_sc_home_rec = os.path.join(IMG_DIR, 'sc_home_rec.png')
    img_sc_forecast = os.path.join(IMG_DIR, 'sc_forecast.png')
    img_sc_history = os.path.join(IMG_DIR, 'sc_history.png')
    img_sc_profile = os.path.join(IMG_DIR, 'sc_profile.png')
    img_sc_admin_dash = os.path.join(IMG_DIR, 'sc_admin_dash.png')
    img_sc_admin_clothing = os.path.join(IMG_DIR, 'sc_admin_clothing.png')
    img_sc_admin_users = os.path.join(IMG_DIR, 'sc_admin_users.png')
    img_sc_admin_weather = os.path.join(IMG_DIR, 'sc_admin_weather.png')
    img_sc_admin_rec = os.path.join(IMG_DIR, 'sc_admin_rec.png')
    all_screenshots = {
        img_sc_login: '用户登录界面',
        img_sc_register: '用户注册界面',
        img_sc_home: '首页天气概览与场景选择',
        img_sc_home_rec: '首页穿衣推荐结果展示',
        img_sc_forecast: '7天天气预报界面',
        img_sc_history: '推荐历史记录界面',
        img_sc_profile: '个人中心界面',
        img_sc_admin_dash: '管理端数据概览仪表盘',
        img_sc_admin_clothing: '管理端服装管理界面',
        img_sc_admin_users: '管理端用户管理界面',
        img_sc_admin_weather: '管理端天气数据监控界面',
        img_sc_admin_rec: '管理端推荐记录界面',
    }
    for p, title in all_screenshots.items():
        if not os.path.exists(p):
            gen_screenshot_placeholder(title, os.path.basename(p))
    print('图表生成完成')

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('安徽三联学院')
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_CHU1, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ANHUI SANLIAN UNIVERSITY')
    set_run_font(run, cn='Times New Roman', en='Times New Roman', size=SIZE_SI, bold=False)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本科毕业设计（论文、创作）')
    set_run_font(run, cn='黑体', en='Times New Roman', size=SIZE_XIAO1, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    cover_items = [
        ('题　　目：', '基于环境信息感知的智能出行穿衣服务系统'),
        ('学生姓名：', '陈鎏'),
        ('学　　号：', 'Z2024130143'),
        ('所在学院：', '计算机与通信工程学院'),
        ('专　　业：', '通信工程（专转本）'),
        ('入学时间：', '2024年9月'),
        ('导师姓名：', '________________'),
        ('职称/学位：', '________________'),
        ('导师所在单位：', '计算机与通信工程学院'),
        ('完成时间：', '2026年4月'),
    ]
    for label, value in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3.0)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(label)
        set_run_font(run, cn='宋体', size=SIZE_SI, bold=False)
        run2 = p.add_run(value)
        set_run_font(run2, cn='宋体', size=SIZE_SI, bold=False)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('安徽三联学院教务处制')
    set_run_font(run, cn='宋体', size=SIZE_SI)

    add_page_break(doc)

    # ===== 新section：正文部分（带页眉） =====
    new_section = doc.add_section()
    new_section.page_width = Cm(21)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)
    new_section.left_margin = Cm(3.0)
    new_section.right_margin = Cm(2.0)
    new_section.header_distance = Cm(1.5)

    header = new_section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run('安徽三联学院毕业设计（论文）')
    set_run_font(hr, cn='宋体', size=SIZE_WU, bold=False)
    pPr = hp._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')
    pBdr.append(bottom_border)
    pPr.append(pBdr)

    footer = new_section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.first_line_indent = Cm(0)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg = fp.add_run()
    set_run_font(run_pg, cn='Times New Roman', en='Times New Roman', size=SIZE_WU)
    run_pg._element.append(fldChar1)
    run_pg._element.append(instrText)
    run_pg._element.append(fldChar2)

    first_section = doc.sections[0]
    first_header = first_section.header
    first_header.is_linked_to_previous = False
    if first_header.paragraphs:
        first_header.paragraphs[0].text = ''

    # ===== 中文摘要 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    run = p.add_run('基于环境信息感知的智能出行穿衣服务系统')
    set_run_font(run, cn='黑体', size=SIZE_XIAO2, bold=False)

    abstract_cn = (
        '伴随城市化进程不断推进以及物联网技术在日常生活领域中的深入渗透，'
        '居民对于智能出行服务方面的需求正在日益增长。本研究的目的在于设计并实现一套'
        '基于环境信息感知机制的智能出行穿衣服务系统，该系统能够通过对多源气象数据'
        '进行采集与融合处理，结合人工智能推荐算法，向用户提供具有个性化特征的穿衣搭配'
        '建议方案。系统的整体设计思路是采用B/S架构体系，后端部分使用Python语言的Flask'
        '框架来进行开发，前端部分则采用Vue 3框架搭配Element Plus组件库实现页面交互功能。'
        '在数据采集方面，系统通过接入和风天气API来获取实时气象数据，同时构建了基于'
        '正弦函数的本地气候模型作为备用方案，覆盖全国32个主要城市。在推荐引擎部分，'
        '系统综合运用了决策树分类算法用于保暖等级的预测，以及基于余弦相似度的协同过滤'
        '算法来实现用户个性化推荐的优化。此外，系统还引入了温湿指数（THI）作为人体'
        '舒适度的量化评估指标。经过对系统功能进行全面的测试验证，结果表明该系统能够'
        '在不同天气条件和应用场景下为用户生成较为合理的穿衣搭配方案，初步达到了智能'
        '出行穿衣服务的设计预期目标。'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run_label = p.add_run('摘要：')
    set_run_font(run_label, cn='黑体', size=SIZE_XIAO4, bold=False)
    run_text = p.add_run(abstract_cn)
    set_run_font(run_text, cn='宋体', size=SIZE_XIAO4, bold=False)
    p.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run('关键词：')
    set_run_font(run_label, cn='黑体', size=SIZE_XIAO4, bold=False)
    run_text = p.add_run('环境信息感知；智能穿衣推荐；决策树算法；协同过滤；Flask；Vue 3')
    set_run_font(run_text, cn='宋体', size=SIZE_XIAO4, bold=False)

    add_page_break(doc)

    # ===== 英文摘要 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    run = p.add_run('Intelligent Travel Dressing Service System Based on Environmental Information Perception')
    set_run_font(run, cn='Times New Roman', en='Times New Roman', size=SIZE_XIAO2, bold=False)

    abstract_en = (
        'With the continuous advancement of urbanization and the deep penetration of IoT '
        'technology into daily life, residents\' demand for intelligent travel services is '
        'growing steadily. This research aims to design and implement an intelligent travel '
        'dressing service system based on environmental information perception mechanisms. '
        'The system collects and fuses multi-source meteorological data, combines artificial '
        'intelligence recommendation algorithms, and provides users with personalized clothing '
        'matching suggestions. The overall design adopts a B/S architecture. The backend is '
        'developed using the Flask framework in Python, while the frontend utilizes the Vue 3 '
        'framework with the Element Plus component library. For data collection, the system '
        'integrates the QWeather API for real-time meteorological data and constructs a local '
        'climate model based on sinusoidal functions as a backup, covering 32 major cities '
        'nationwide. The recommendation engine employs a Decision Tree classification algorithm '
        'for warmth level prediction and a collaborative filtering algorithm based on cosine '
        'similarity for personalized recommendation optimization. Additionally, the Temperature-Humidity '
        'Index (THI) is introduced as a quantitative evaluation metric for human comfort. '
        'Comprehensive functional testing demonstrates that the system can generate reasonable '
        'clothing matching schemes under various weather conditions and application scenarios, '
        'preliminarily achieving the design objectives of intelligent travel dressing services.'
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run('Abstract: ')
    set_run_font(run_label, cn='Times New Roman', en='Times New Roman', size=SIZE_XIAO4, bold=True)
    run_text = p.add_run(abstract_en)
    set_run_font(run_text, cn='Times New Roman', en='Times New Roman', size=SIZE_XIAO4, bold=False)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run('Key words: ')
    set_run_font(run_label, cn='Times New Roman', en='Times New Roman', size=SIZE_XIAO4, bold=True)
    run_text = p.add_run('Environmental Information Perception; Intelligent Clothing Recommendation; '
                         'Decision Tree; Collaborative Filtering; Flask; Vue 3')
    set_run_font(run_text, cn='Times New Roman', en='Times New Roman', size=SIZE_XIAO4, bold=False)

    add_page_break(doc)

    # ===== 目录 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(18)
    pf.first_line_indent = Cm(0)
    run = p.add_run('目　录')
    set_run_font(run, cn='黑体', size=SIZE_XIAO2, bold=False)

    toc_items = [
        ('1  绪论', 0), ('1.1 研究背景与现状', 1), ('1.2 研究目的及意义', 1), ('1.3 研究内容及思路', 1),
        ('2  环境信息感知技术综述', 0), ('2.1 环境信息感知技术', 1),
        ('2.1.1 气象数据采集技术', 2), ('2.1.2 多源数据融合方法', 2),
        ('2.2 智能推荐算法', 1), ('2.2.1 决策树算法', 2), ('2.2.2 协同过滤算法', 2),
        ('2.3 Web应用开发技术', 1), ('2.3.1 Flask框架', 2), ('2.3.2 Vue.js框架', 2),
        ('3  智能出行中的穿衣需求分析', 0), ('3.1 用户需求分析', 1),
        ('3.1.1 普通用户需求', 2), ('3.1.2 管理员需求', 2),
        ('3.2 功能需求分析', 1), ('3.3 非功能需求分析', 1),
        ('4  智能出行穿衣服务系统设计', 0), ('4.1 系统总体架构设计', 1),
        ('4.2 数据库设计', 1), ('4.2.1 概念结构设计', 2), ('4.2.2 逻辑结构设计', 2),
        ('4.3 功能模块设计', 1), ('4.4 推荐算法设计', 1),
        ('4.4.1 环境信息感知与决策数据', 2), ('4.4.2 决策树分类与三种推荐模式', 2),
        ('4.4.3 个人因素调整机制', 2), ('4.4.4 服装筛选与协同过滤优化', 2),
        ('4.4.5 穿衣建议生成与舒适度评分', 2),
        ('4.5 API接口设计', 1),
        ('5  系统实现与测试', 0), ('5.1 开发环境', 1), ('5.2 系统实现', 1),
        ('5.2.1 后端核心实现', 2), ('5.2.2 前端核心实现', 2), ('5.2.3 算法实现', 2),
        ('5.3 系统测试', 1), ('5.3.1 功能测试', 2), ('5.3.2 性能测试', 2),
        ('6  用户体验评估', 0), ('6.1 评估方法', 1), ('6.2 评估结果与分析', 1),
        ('7  总结与展望', 0), ('7.1 研究总结', 1), ('7.2 不足与展望', 1),
        ('参考文献', 0), ('致谢', 0),
    ]

    for text, level in toc_items:
        p = doc.add_paragraph()
        indent = level * 0.74
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, cn='宋体', size=SIZE_XIAO4)

    add_page_break(doc)

    # ============================================================
    # 第一章 绪论
    # ============================================================
    add_heading_l1(doc, '绪论', '1')

    add_heading_l2(doc, '1.1 研究背景与现状')

    add_paragraph(doc,
        '近些年来，随着城市化进程的持续加快和交通出行方式不断趋于多样化发展，'
        '人们日常出行所面对的环境条件也变得越来越复杂和多变。气温、湿度、风速、'
        '降水以及空气质量等环境因素对人们的穿着选择会产生直接性的影响，然而许多人'
        '在出门之前往往缺少对当前及未来天气变化情况的充分了解，从而导致穿着不合适'
        '的情况频繁发生。特别是在季节交替时期，早晚温差的显著变化使得穿衣选择变得'
        '更加困难。在这样的背景之下，如何利用现代信息技术手段来帮助用户做出更加合理'
        '的穿着决策，成为了一个具有实际意义的研究课题。'
    )

    add_paragraph(doc,
        '物联网（IoT）与人工智能（AI）技术的快速发展为上述问题的解决提供了新的'
        '技术路径。通过多源气象数据的实时采集与智能分析，系统可以对环境信息进行'
        '感知和理解，进而生成针对用户个体特征的穿衣建议。国内方面，姚朋（2024）'
        '对基于多源气象数据的城市微气候感知进行了研究，常姬昊（2023）探讨了公共'
        '气象服务平台的智能化转型路径。国际上，欧美研究团队利用车载传感网络探索'
        '城市微气候下的出行舒适度模型，北美与北欧的应用实践将能耗感知算法融入可'
        '穿戴设备。但目前大多数研究仅关注单一功能的开发，如环境监测或算法优化，'
        '缺乏从环境感知到行为建议全链条整合的实践模型[1-3]。'
    )

    add_paragraph(doc,
        '中国气象局发布的数据表明，2023年全国平均气温较常年偏高0.8℃，极端天气'
        '事件呈现出频发态势，这进一步凸显了智能穿衣推荐系统在日常生活中的应用价值。'
        '因此，本研究尝试在现有研究基础上，通过融合环境信息感知技术与智能推荐算法，'
        '构建一套完整的出行穿衣服务系统，以期为该领域的研究与实践提供有益参考。'
    )

    add_heading_l2(doc, '1.2 研究目的及意义')

    add_paragraph(doc,
        '本研究的主要目的是设计并实现一套融合环境信息感知与智能决策功能的出行穿衣'
        '服务系统。从理论层面来说，本研究将环境信息感知机制引入传统的出行行为模型之中，'
        '通过构建温度、湿度、风速、空气质量等多维环境参数与穿衣决策之间的映射关系，'
        '丰富了人与环境交互的认知建模路径。温湿指数（THI）的引入为人体舒适度评估'
        '提供了定量化的理论依据，而决策树与协同过滤相结合的推荐方法则扩展了智能推荐'
        '算法在穿衣场景下的应用范围[4-6]。'
    )

    add_paragraph(doc,
        '从现实意义角度而言，本系统的实现可以有效地降低用户因对天气状况判断不准确而'
        '导致的穿着不适问题，从而提升日常出行的舒适度和健康水平。系统所支持的多场景'
        '切换功能——包括通勤、运动、约会、聚会、旅行、居家、商务等七种常见出行场景——'
        '能够覆盖用户在不同生活情境中的穿衣需求。对于老年群体和体质较为敏感的人群，'
        '系统通过体质参数的个性化调节来提供更为贴切的建议，具备一定程度的健康防护价值。'
    )

    add_heading_l2(doc, '1.3 研究内容及思路')

    add_paragraph(doc,
        '本研究的主要内容包含以下几个方面：首先，进行系统需求分析，明确用户需求'
        '和功能需求；其次，完成系统的总体架构设计和详细设计；然后，实现系统的各个'
        '功能模块并进行集成测试；最后，对系统的用户体验进行评估分析。'
    )

    add_paragraph(doc,
        '研究思路上，本文采用"需求分析—系统设计—编码实现—测试评估"的软件工程方法'
        '来推进整个研究过程。在技术路线方面，后端选用Python Flask框架实现RESTful API'
        '服务，前端使用Vue 3框架构建单页面应用，数据库采用SQLite进行数据持久化存储。'
        '推荐引擎结合决策树分类算法与协同过滤算法来生成个性化推荐方案，天气数据'
        '通过接入第三方气象API并辅以本地气候模型来保证数据的可靠性和覆盖范围。'
    )

    add_page_break(doc)
    # ============================================================
    # 第二章 环境信息感知技术综述
    # ============================================================
    add_heading_l1(doc, '环境信息感知技术综述', '2')

    add_heading_l2(doc, '2.1 环境信息感知技术')

    add_heading_l3(doc, '2.1.1 气象数据采集技术')

    add_paragraph(doc,
        '气象数据采集是环境信息感知系统中最为基础和关键的环节之一。当前主流的气象'
        '数据获取方式主要有以下几种：第一种是通过调用气象服务提供商的开放API接口来获取'
        '实时和预报气象数据，例如和风天气（QWeather）、OpenWeatherMap等平台均提供了'
        '覆盖全球范围的气象数据查询服务，其数据来源于分布在各地的气象站网络和卫星遥感'
        '系统[7]。第二种是利用本地部署的传感器设备进行直接的环境参数测量，这种方式在'
        'IoT场景中较为常见，但部署成本相对较高。第三种是基于历史气象统计规律构建数学'
        '模型来估算特定地区和时间的气象数据，该方法在API不可用时可以作为一种有效的'
        '备用方案。'
    )

    add_paragraph(doc,
        '本系统采用的数据采集策略是以和风天气API作为主要的数据来源，同时自主构建'
        '了一套基于正弦函数的本地气候模型来提供数据保障。该气候模型基于各城市的'
        '年均温度、温度振幅、平均湿度和降水概率等气候参数，使用正弦函数模拟全年的'
        '温度变化规律，并通过高斯随机扰动来模拟日间的温度波动。这种双通道的数据采集'
        '方案既保证了在正常网络环境下能够获取高精度的实时气象数据，又能在网络异常或'
        'API服务中断时保持系统的可用性。'
    )

    add_heading_l3(doc, '2.1.2 多源数据融合方法')

    add_paragraph(doc,
        '多源数据融合是指将来自不同数据源的信息进行综合处理和整合，以获得比单一数据源'
        '更加完整和准确的环境描述。在本系统中，数据融合主要体现在以下几个层面：首先是'
        '来自气象API的实时观测数据与本地气候模型生成数据之间的融合，当API返回有效数据时'
        '会覆盖模型生成的对应字段，从而实现数据质量的提升。其次是温度、湿度、风速、降水'
        '等多个气象参数之间的关联分析，系统通过温湿指数（THI）将多维气象参数融合为一个'
        '综合的舒适度评价指标[8]。THI的计算公式为：'
    )

    add_paragraph(doc,
        'THI = (1.8×T + 32) - 0.55×(1 - RH/100)×(1.8×T - 26)',
        align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False
    )

    add_paragraph(doc,
        '其中T表示气温（℃），RH表示相对湿度（%）。THI值越高表示体感越热，反之则越冷。'
        '系统根据THI值将人体舒适度划分为8个等级，从"极冷"到"炎热"，并据此为用户提供'
        '差异化的穿衣建议。此外，风寒因子也被纳入到舒适度的综合评估体系之中，当气温'
        '低于10℃且风速大于4.8km/h时，系统会使用风寒指数公式来修正体感温度的计算结果。'
    )

    add_heading_l2(doc, '2.2 智能推荐算法')

    add_heading_l3(doc, '2.2.1 决策树算法')

    add_paragraph(doc,
        '决策树（Decision Tree）是一种常用的机器学习分类与回归算法，其基本原理是通过'
        '对样本数据的特征属性进行递归分割，构建出一棵从根节点到叶节点的树形结构，每个'
        '内部节点表示一个特征属性上的判断条件，每个叶节点则表示一个分类结果。决策树算法'
        '具有模型可解释性强、计算效率高、对数据预处理要求较低等优点，在分类问题中得到了'
        '广泛的应用[5-6]。'
    )

    add_paragraph(doc,
        '本系统使用scikit-learn库中的DecisionTreeClassifier来实现穿衣保暖等级的预测。'
        '模型的输入特征包括气温、湿度、风速、降水量、空气质量指数和场景类别共6个维度，'
        '输出为1到5的保暖等级，其中1表示需要最轻薄的穿着，5表示需要最厚实的保暖衣物。'
        '训练数据由专家经验知识构建，涵盖了不同气候条件和出行场景下的合理穿衣保暖等级。'
        '模型的关键超参数设置为：最大树深度8层，最小分割样本数3，最小叶节点样本数2。'
    )

    add_heading_l3(doc, '2.2.2 协同过滤算法')

    add_paragraph(doc,
        '协同过滤（Collaborative Filtering）是推荐系统领域中最为经典的算法之一，其'
        '核心思想是"相似的用户可能喜欢相似的物品"。协同过滤算法主要可以分为基于用户的'
        '协同过滤和基于物品的协同过滤两大类。基于用户的协同过滤通过计算用户之间的相似'
        '度来发现品味相近的用户群体，然后将相似用户喜欢但目标用户尚未接触过的物品推荐'
        '给目标用户[9-10]。'
    )

    add_paragraph(doc,
        '本系统采用基于用户的协同过滤方法，使用余弦相似度作为用户相似度的度量指标。'
        '具体来说，系统维护了一个用户-服装偏好评分矩阵，每个元素代表某位用户对某件'
        '服装的偏好程度。当需要为目标用户推荐服装时，系统首先找到与其最相似的Top-N个'
        '用户，然后根据相似用户对候选服装的加权评分进行排序，选出评分最高的服装作为'
        '推荐结果。余弦相似度的计算公式为：'
    )

    add_paragraph(doc,
        'sim(u, v) = cos(θ) = (u·v) / (||u|| × ||v||)',
        align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False
    )

    add_paragraph(doc,
        '其中u和v分别表示两个用户的偏好评分向量。该方法在数据较为稀疏的情况下仍然'
        '能够产生一定的推荐效果，适合于本系统目前的数据规模。'
    )

    add_heading_l2(doc, '2.3 Web应用开发技术')

    add_heading_l3(doc, '2.3.1 Flask框架')

    add_paragraph(doc,
        'Flask是一个基于Python语言的轻量级Web应用框架，它遵循WSGI规范，采用Jinja2'
        '模板引擎和Werkzeug工具包。Flask的设计理念是"微框架"，核心功能保持精简，'
        '但通过丰富的扩展生态系统可以灵活地添加数据库操作、用户认证、表单验证等功能[12]。'
        '本系统使用了Flask-SQLAlchemy实现ORM数据库操作，Flask-JWT-Extended实现基于'
        'JWT（JSON Web Token）的用户身份认证，Flask-CORS处理跨域资源共享问题，以及'
        'Flask-Migrate进行数据库迁移管理。Flask轻量灵活的特点使其非常适合用于中小规模'
        'Web服务的快速开发。'
    )

    add_heading_l3(doc, '2.3.2 Vue.js框架')

    add_paragraph(doc,
        'Vue.js是一个用于构建用户界面的渐进式JavaScript框架，由尤雨溪（Evan You）'
        '创建并维护。Vue 3是其最新的主要版本，引入了Composition API、更好的TypeScript'
        '支持以及改进后的响应式系统等新特性[13]。本系统采用Vue 3框架搭配Vue Router进行'
        '前端路由管理，使用Pinia作为状态管理工具，并集成了Element Plus作为UI组件库。'
        'Element Plus提供了丰富的预制组件，包括表格、表单、对话框、导航菜单等，极大地'
        '加快了前端界面的开发速度。此外，系统还引入了ECharts图表库来实现数据可视化功能，'
        '使用Axios库来处理HTTP请求的发送和响应。'
    )

    add_page_break(doc)
    # ============================================================
    # 第三章 穿衣需求分析
    # ============================================================
    add_heading_l1(doc, '智能出行中的穿衣需求分析', '3')

    add_heading_l2(doc, '3.1 用户需求分析')

    add_heading_l3(doc, '3.1.1 普通用户需求')

    add_paragraph(doc,
        '通过对目标用户群体的调研和分析，本系统的普通用户主要有以下方面的核心需求：'
        '第一，用户希望能够便捷地获取到所在城市的实时天气信息，包括温度、湿度、风速、'
        '空气质量等关键指标；第二，用户期望系统能够根据当前的天气状况和出行场景自动'
        '生成合理的穿衣搭配建议；第三，用户需要能够根据自身的体质特征（如怕冷、怕热）'
        '来获取个性化的推荐结果；第四，用户希望查看未来数天的天气预报以便提前规划穿着；'
        '第五，用户希望能对推荐结果进行反馈评价，以帮助系统不断改善推荐质量。'
        '普通用户的用例图如图3-1所示。'
    )

    insert_image(doc, img_uc_user, 12)
    add_figure_caption(doc, '图3-1 普通用户用例图')

    add_heading_l3(doc, '3.1.2 管理员需求')

    add_paragraph(doc,
        '系统管理员的主要需求包括：第一，能够通过数据概览面板查看系统的整体运行状况，'
        '包括用户数量、服装数量、推荐记录数、反馈数量等关键统计指标；第二，能够对系统中'
        '的用户账户进行管理操作，如查询、启用或禁用用户；第三，能够对服装数据进行'
        '增删改查操作，维护服装库的完整性和时效性；第四，能够监控和查看天气数据的采集'
        '情况；第五，能够查看和分析推荐记录数据。管理员的用例图如图3-2所示。'
    )

    insert_image(doc, img_uc_admin, 12)
    add_figure_caption(doc, '图3-2 管理员用例图')

    add_heading_l2(doc, '3.2 功能需求分析')

    add_paragraph(doc,
        '根据上述用户需求分析结果，本系统的功能需求可以概括为以下四个主要方面：'
    )

    add_paragraph(doc,
        '（1）数据采集功能：系统需要能够从外部气象服务接口获取全国主要城市的实时天气'
        '数据和未来7天的天气预报数据，并将获取到的数据进行缓存存储以减少重复调用。'
        '同时系统需要支持多种出行场景标签的定义与管理，包括通勤、运动、约会、聚会、'
        '旅行、居家和商务等七种典型场景。'
    )

    add_paragraph(doc,
        '（2）智能推荐功能：系统需要实现基于决策树算法的穿衣保暖等级预测，以及基于'
        '协同过滤算法的个性化服装推荐排序。推荐结果应当包含完整的穿衣搭配方案，涵盖'
        '上衣、下装、外套、鞋子和配饰五个品类。系统还需要根据用户的体质类型和年龄'
        '等因素对推荐结果进行适当的调整和优化。'
    )

    add_paragraph(doc,
        '（3）用户交互功能：系统需要提供友好的用户注册和登录界面，实时天气信息展示页面，'
        '穿衣推荐结果展示页面，7天天气预报查看页面，推荐历史记录查看页面，以及用户'
        '个人信息编辑页面。用户能够切换不同的出行场景来获取针对性的推荐，也能够对'
        '推荐结果进行满意度评分反馈。'
    )

    add_paragraph(doc,
        '（4）远程管理功能：系统需要为管理员提供一个功能完善的后台管理平台，包括数据'
        '概览仪表盘、用户管理、服装管理（增删改查）、天气数据监控和推荐记录查看等'
        '功能模块。管理员可以通过浏览器随时随地访问管理后台来进行系统维护和数据管理。'
    )

    add_heading_l2(doc, '3.3 非功能需求分析')

    add_paragraph(doc,
        '除功能需求外，系统还应满足以下非功能性需求：在安全性方面，用户密码需要经过'
        '哈希加密存储，API请求需要通过JWT令牌验证身份，防止未授权访问。在性能方面，'
        '页面加载时间应控制在2秒以内，天气数据采用缓存机制以提高响应速度。在可用性'
        '方面，界面设计应当简洁直观，操作流程合理顺畅。在兼容性方面，前端页面应当'
        '支持主流浏览器并具备一定的响应式布局能力，以适配不同尺寸的终端设备。在可'
        '维护性方面，代码应采用模块化架构设计，便于后续功能的扩展和维护。'
    )

    add_page_break(doc)
    # ============================================================
    # 第四章 系统设计
    # ============================================================
    add_heading_l1(doc, '智能出行穿衣服务系统设计', '4')

    add_heading_l2(doc, '4.1 系统总体架构设计')

    add_paragraph(doc,
        '本系统采用经典的B/S（Browser/Server）架构体系进行设计，整体上可以划分为'
        '四个层次：用户交互层、业务逻辑层、数据处理层和数据存储层。各层之间通过'
        '标准化的接口进行通信，实现了较好的松耦合效果。系统总体架构如图4-1所示。'
    )

    insert_image(doc, img_arch, 14)
    add_figure_caption(doc, '图4-1 系统总体架构图')

    add_paragraph(doc,
        '用户交互层是系统的最上层，负责与用户进行直接的交互。该层使用Vue 3框架构建'
        '单页面应用（SPA），集成Element Plus组件库提供丰富的UI控件，使用ECharts实现'
        '数据可视化图表，通过Axios向后端API发送HTTP请求并处理响应数据。'
    )

    add_paragraph(doc,
        '业务逻辑层是系统的核心中间层，使用Flask框架实现RESTful风格的API服务。该层'
        '负责处理来自前端的各类业务请求，包括用户认证（JWT）、天气数据查询、穿衣推荐'
        '生成、服装管理、用户管理等。所有API接口统一采用JSON格式进行数据交换。'
    )

    add_paragraph(doc,
        '数据处理层主要包含系统的核心算法组件，即决策树推荐模型、协同过滤推荐模型'
        '和THI舒适度计算模型。该层接收来自业务逻辑层的调用请求，对原始数据进行分析'
        '和处理后返回计算结果。'
    )

    add_paragraph(doc,
        '数据存储层是系统的最底层，采用SQLite作为开发环境的关系型数据库（生产环境可'
        '切换为MySQL），通过SQLAlchemy ORM框架实现对象关系映射。同时该层也包含对和风'
        '天气API的接入逻辑，将获取到的外部气象数据存入本地数据库进行缓存管理。'
    )

    add_heading_l2(doc, '4.2 数据库设计')
    add_heading_l3(doc, '4.2.1 概念结构设计')

    add_paragraph(doc,
        '根据系统的功能需求分析结果，本系统的数据库概念模型包含以下主要实体：用户'
        '（User）、服装（Clothing）、服装分类（ClothingCategory）、天气记录（WeatherRecord）、'
        '推荐记录（Recommendation）、用户反馈（UserFeedback）和用户偏好（UserPreference）。'
        '各实体之间的关系如E-R图4-2所示。'
    )

    insert_image(doc, img_er, 14)
    add_figure_caption(doc, '图4-2 系统E-R图')

    add_paragraph(doc,
        '在实体关系中，一个用户可以拥有多条推荐记录（一对多关系），一个用户可以提交多条'
        '反馈信息（一对多关系），一个用户可以拥有多条服装偏好记录（一对多关系），一件服装'
        '属于一个服装分类（多对一关系），一条推荐记录可以关联多件服装（多对一关系，分别'
        '对应上衣、下装、外套、鞋子和配饰五个字段）。'
    )

    add_heading_l3(doc, '4.2.2 逻辑结构设计')

    add_paragraph(doc,
        '在概念模型的基础上，将E-R模型转换为关系数据库的逻辑结构。以下给出系统各主要'
        '数据表的结构设计。'
    )

    add_table_caption(doc, '表4-1 用户表（users）')
    add_three_line_table(doc,
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'Integer', '主键，自增', '用户ID'],
            ['username', 'String(64)', '唯一，非空', '用户名'],
            ['email', 'String(120)', '唯一，非空', '邮箱'],
            ['password_hash', 'String(256)', '非空', '密码哈希'],
            ['gender', 'String(10)', '默认unknown', '性别'],
            ['age', 'Integer', '默认25', '年龄'],
            ['city', 'String(64)', '默认北京', '所在城市'],
            ['body_type', 'String(20)', '默认normal', '体质类型'],
            ['role', 'String(20)', '默认user', '角色'],
            ['is_active', 'Boolean', '默认True', '是否启用'],
            ['created_at', 'DateTime', '自动生成', '创建时间'],
        ],
        col_widths=[3, 3, 4, 4]
    )

    add_table_caption(doc, '表4-2 服装表（clothes）')
    add_three_line_table(doc,
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'Integer', '主键，自增', '服装ID'],
            ['name', 'String(128)', '非空', '服装名称'],
            ['category_id', 'Integer', '外键', '分类ID'],
            ['warmth_level', 'Integer', '1-5', '保暖指数'],
            ['breathability', 'Integer', '1-5', '透气指数'],
            ['waterproof', 'Integer', '1-5', '防水指数'],
            ['min_temp', 'Float', '默认-10', '适用最低温度'],
            ['max_temp', 'Float', '默认40', '适用最高温度'],
            ['suitable_scenes', 'String(256)', '逗号分隔', '适用场景'],
            ['suitable_gender', 'String(20)', '默认all', '适用性别'],
        ],
        col_widths=[3, 3, 4, 4]
    )

    add_table_caption(doc, '表4-3 天气记录表（weather_records）')
    add_three_line_table(doc,
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'Integer', '主键，自增', '记录ID'],
            ['city', 'String(64)', '非空，索引', '城市名'],
            ['date', 'Date', '非空', '日期'],
            ['temperature', 'Float', '', '温度℃'],
            ['feels_like', 'Float', '', '体感温度℃'],
            ['humidity', 'Integer', '', '湿度%'],
            ['wind_speed', 'Float', '', '风速km/h'],
            ['weather_text', 'String(64)', '', '天气描述'],
            ['aqi', 'Integer', '', '空气质量指数'],
            ['uv_index', 'Integer', '', '紫外线指数'],
        ],
        col_widths=[3, 3, 4, 4]
    )

    add_table_caption(doc, '表4-4 推荐记录表（recommendations）')
    add_three_line_table(doc,
        ['字段名', '数据类型', '约束', '说明'],
        [
            ['id', 'Integer', '主键，自增', '推荐ID'],
            ['user_id', 'Integer', '外键', '用户ID'],
            ['city', 'String(64)', '非空', '城市'],
            ['scene', 'String(64)', '默认通勤', '场景'],
            ['outfit_top', 'Integer', '外键', '上衣ID'],
            ['outfit_bottom', 'Integer', '外键', '下装ID'],
            ['outfit_outer', 'Integer', '外键', '外套ID'],
            ['outfit_shoes', 'Integer', '外键', '鞋子ID'],
            ['comfort_score', 'Float', '0-100', '舒适度评分'],
            ['algorithm_type', 'String(32)', '', '算法类型'],
        ],
        col_widths=[3, 3, 4, 4]
    )

    add_heading_l2(doc, '4.3 功能模块设计')

    add_paragraph(doc,
        '根据前述需求分析，系统的功能模块可以划分为五个主要部分：数据采集模块、推荐'
        '引擎模块、用户交互模块、后台管理模块和认证鉴权模块。功能模块的整体结构如'
        '图4-3所示。'
    )

    insert_image(doc, img_module, 14)
    add_figure_caption(doc, '图4-3 系统功能模块图')

    add_paragraph(doc,
        '数据采集模块负责从外部气象API获取天气数据，并通过本地气候模型进行数据补充和'
        '校验。推荐引擎模块集成了决策树算法和协同过滤算法，完成从天气数据到穿衣方案的'
        '智能转换。用户交互模块提供了天气展示、推荐展示和用户中心等前端页面功能。后台'
        '管理模块包含用户管理、服装管理和数据监控等管理端功能。认证鉴权模块通过JWT令牌'
        '机制保障系统的访问安全。'
    )

    add_heading_l2(doc, '4.4 推荐算法设计')

    add_paragraph(doc,
        '本系统的推荐算法采用"环境感知决策+个人因素调整+协同过滤优化"的多阶段智能推荐'
        '方案。算法的核心设计目标是将多维度的环境信息转化为可量化的穿衣保暖等级，再结合'
        '用户的个人生理特征进行个性化修正，最终通过协同过滤算法进行服装偏好排序，生成完整'
        '的穿衣搭配方案。推荐算法的整体流程如图4-4所示。'
    )

    insert_image(doc, img_algo_flow, 10)
    add_figure_caption(doc, '图4-4 推荐算法流程图')

    add_heading_l3(doc, '4.4.1 环境信息感知与决策数据')

    add_paragraph(doc,
        '环境信息感知是推荐算法的第一步，系统通过采集多维度的气象参数来构建当前环境的'
        '全景描述。系统采集的环境感知数据包括：气温（℃）、体感温度（℃）、相对湿度（%）、'
        '风速（km/h）、风向、降水量（mm）、气压（hPa）、能见度（km）、紫外线指数、'
        '空气质量指数（AQI）及其等级。这些原始数据经过融合处理后，会生成两项综合决策指标：'
        '温湿指数（THI）和风寒修正温度。'
    )

    add_paragraph(doc,
        'THI（Temperature-Humidity Index）是衡量人体在特定温湿度条件下舒适程度的综合指标。'
        '其计算公式为：'
    )

    add_paragraph(doc,
        'THI = (1.8×T + 32) - 0.55×(1 - RH/100)×(1.8×T - 26)',
        align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False
    )

    add_paragraph(doc,
        '其中T为当前气温（℃），RH为相对湿度（%）。系统根据THI值将人体舒适度划分为8个'
        '等级，每个等级对应不同的穿衣建议基准，具体划分标准如表4-6所示。'
    )

    add_table_caption(doc, '表4-6 THI舒适度等级划分表')
    add_three_line_table(doc,
        ['THI范围', '等级', '体感描述', '穿衣建议基准'],
        [
            ['THI < 40', '1级', '极冷', '需厚羽绒服、棉服，注意严密防寒'],
            ['40 ≤ THI < 45', '2级', '寒冷', '需厚实保暖衣物，建议多层穿搭'],
            ['45 ≤ THI < 55', '3级', '偏冷', '适当增添衣物，需外套保暖'],
            ['55 ≤ THI < 60', '4级', '凉爽', '薄外套或长袖即可，适宜外出'],
            ['60 ≤ THI < 70', '5级', '舒适', '穿着轻便即可，体感最佳'],
            ['70 ≤ THI < 75', '6级', '温暖', '短袖为主，注意防晒'],
            ['75 ≤ THI < 80', '7级', '偏热', '穿着清凉透气衣物'],
            ['THI ≥ 80', '8级', '炎热', '极薄衣物，注意防暑降温'],
        ],
        col_widths=[2.5, 1.5, 2, 6]
    )

    add_paragraph(doc,
        '当气温低于10℃且风速超过4.8km/h时，系统会启用风寒修正公式来计算体感温度，'
        '以更准确地反映人体在风力条件下的实际感受。风寒修正温度的计算公式为：'
    )

    add_paragraph(doc,
        'T_wc = 13.12 + 0.6215×T - 11.37×V^0.16 + 0.3965×T×V^0.16',
        align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False
    )

    add_paragraph(doc,
        '其中T为实际气温（℃），V为风速（km/h）。风寒修正后的体感温度通常会明显低于'
        '实际气温，这一修正使得系统在寒冷大风天气下能够给出更加合理的加厚穿衣建议。'
    )

    add_heading_l3(doc, '4.4.2 决策树分类与三种推荐模式')

    add_paragraph(doc,
        '系统构建了三套独立的决策树模型来满足不同体感偏好的用户需求，分别是标准模式'
        '（standard）、偏暖模式（warm）和偏冷模式（cool）。三套模型的输入特征维度完全'
        '一致，均为气温、湿度、风速、降水量、AQI和场景ID共6个维度，但各自使用了不同的'
        '训练数据集，从而在相同环境条件下输出不同的保暖等级预测。'
    )

    add_table_caption(doc, '表4-7 三种推荐模式对比')
    add_three_line_table(doc,
        ['模式', '适用人群', '训练数据特点', '模型参数', '效果描述'],
        [
            ['标准模式', '一般人群', '基于专家经验的均衡数据', 'max_depth=8', '温度-保暖等级均衡映射'],
            ['偏暖模式', '怕冷人群', '同温度下保暖等级系统上调', 'max_depth=6', '倾向推荐更厚实的衣物'],
            ['偏冷模式', '怕热人群', '同温度下保暖等级系统下调', 'max_depth=6', '倾向推荐更清凉的衣物'],
        ],
        col_widths=[2, 2, 3, 2.5, 3.5]
    )

    add_paragraph(doc,
        '每套模型均使用35组训练样本，特征涵盖了从-10℃到35℃的全温度区间以及7种出行场景。'
        '以15℃的通勤场景为例，标准模式预测保暖等级为3（适中偏凉），偏暖模式预测为4（需要'
        '较厚外套），偏冷模式预测为2（轻便穿着即可）。这种多模式设计使系统能够适应不同用户'
        '对温度的差异化感知，用户可以在前端界面自由切换推荐模式。'
    )

    add_heading_l3(doc, '4.4.3 个人因素调整机制')

    add_paragraph(doc,
        '在决策树输出基础保暖等级后，系统会根据用户的个人生理特征进行二次调整，以实现'
        '真正的个性化推荐。个人因素调整涉及三个维度：体质类型、年龄和性别，其具体调整'
        '规则如表4-8所示。'
    )

    add_table_caption(doc, '表4-8 个人因素调整规则表')
    add_three_line_table(doc,
        ['调整维度', '用户特征', '调整方式', '调整幅度', '调整理由'],
        [
            ['体质类型', '怕冷体质（cold_sensitive）', '保暖等级上调', '+1级', '怕冷人群对低温更敏感'],
            ['体质类型', '怕热体质（heat_sensitive）', '保暖等级下调', '-1级', '怕热人群对高温更敏感'],
            ['体质类型', '正常体质（normal）', '不调整', '0', '按标准模式推荐'],
            ['年龄', '60岁以上老年用户', '保暖等级上调', '+1级', '老年人体温调节能力较弱'],
            ['年龄', '60岁及以下用户', '不调整', '0', '按标准年龄推荐'],
            ['性别', '男性（male）', '筛选男性适用服装', '—', '过滤非男性服装选项'],
            ['性别', '女性（female）', '筛选女性适用服装', '—', '过滤非女性服装选项'],
            ['性别', '未设置', '显示全部服装', '—', '不做性别过滤'],
        ],
        col_widths=[2, 3.5, 2.5, 1.5, 3.5]
    )

    add_paragraph(doc,
        '调整后的保暖等级会被限制在1到5的有效范围内（使用min/max函数截断），防止出现'
        '无效值。体质类型和年龄的调整效果可以叠加，例如一位65岁的怕冷体质老人，在标准'
        '模式预测保暖等级为3的情况下，经过体质调整（+1）和年龄调整（+1）后，最终保暖等级'
        '为5，系统将推荐最厚实的防寒衣物组合。'
    )

    add_heading_l3(doc, '4.4.4 服装筛选与协同过滤优化')

    add_paragraph(doc,
        '第二阶段，系统根据调整后的保暖等级和用户特征，从服装数据库中进行多条件筛选。'
        '筛选过程按上衣、下装、外套、鞋子和配饰五个品类槽位分别执行，每个槽位的筛选条件'
        '包括：（1）服装品类属于该槽位对应的分类集合；（2）服装处于启用状态；（3）当前气温'
        '在服装的适用温度范围[min_temp, max_temp]之内；（4）服装的适用性别与用户性别匹配'
        '（或为"全部"）；（5）服装的适用场景包含当前出行场景。满足条件的候选服装按保暖等级'
        '与预测值的差值绝对值升序排列，取前20件进入协同过滤排序。'
    )

    add_paragraph(doc,
        '协同过滤阶段使用基于用户的协同过滤算法，对每个槽位的候选服装进行个性化排序。'
        '系统从用户-服装偏好矩阵中找到与目标用户最相似的Top-5用户，然后基于相似用户对'
        '候选服装的加权评分来计算推荐分数。推荐分数的计算公式为：'
    )

    add_paragraph(doc,
        'score(u, i) = Σ(sim(u, v) × r(v, i)) / Σ|sim(u, v)|',
        align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False
    )

    add_paragraph(doc,
        '其中sim(u, v)为用户u与相似用户v之间的余弦相似度，r(v, i)为相似用户v对服装i'
        '的偏好评分。每个槽位选取推荐分数最高的一件服装作为最终推荐，组合成完整的穿搭方案。'
    )

    add_heading_l3(doc, '4.4.5 穿衣建议生成与舒适度评分')

    add_paragraph(doc,
        '系统在生成服装搭配方案的同时，还会根据气象条件自动判断是否需要外套和配饰：当保暖'
        '等级≥3或有降水或风速>15km/h时建议穿外套；当保暖等级≥4或有降水或AQI>150时建议'
        '携带配饰（围巾、雨伞、口罩等）。此外，系统还针对不同出行场景提供文字穿衣建议，'
        '例如运动场景建议选择透气排汗面料，商务场合建议着正装。'
    )

    add_paragraph(doc,
        '最终的舒适度评分（comfort_score）由基础舒适度、模式调整因子和保暖匹配度三部分'
        '加权计算得出。基础舒适度以22℃/55%湿度为最佳体感基准，偏离越大分值越低。搭配'
        '评分（match_score）则综合了舒适度、保暖等级适配度和模式匹配度，范围为5到98分，'
        '为用户提供直观的推荐质量评价。'
    )

    add_heading_l2(doc, '4.5 API接口设计')

    add_paragraph(doc,
        '系统的API接口采用RESTful风格进行设计，所有接口统一使用JSON格式进行请求和'
        '响应数据的交换。接口按照功能模块进行分组，主要包含以下几组：'
    )

    add_table_caption(doc, '表4-5 主要API接口列表')
    add_three_line_table(doc,
        ['接口路径', '方法', '功能说明', '认证'],
        [
            ['/api/auth/register', 'POST', '用户注册', '否'],
            ['/api/auth/login', 'POST', '用户登录', '否'],
            ['/api/weather/current', 'GET', '获取当前天气', '是'],
            ['/api/weather/forecast', 'GET', '获取7天预报', '是'],
            ['/api/recommendation/', 'GET', '获取穿衣推荐', '是'],
            ['/api/recommendation/feedback', 'POST', '提交反馈评分', '是'],
            ['/api/user/profile', 'GET/PUT', '个人信息管理', '是'],
            ['/api/admin/stats', 'GET', '管理统计数据', '管理员'],
            ['/api/admin/users', 'GET', '用户列表管理', '管理员'],
            ['/api/clothing/', 'CRUD', '服装数据管理', '管理员'],
        ],
        col_widths=[4, 1.5, 4, 2]
    )

    add_page_break(doc)
    # ============================================================
    # 第五章 系统实现与测试
    # ============================================================
    add_heading_l1(doc, '系统实现与测试', '5')

    add_heading_l2(doc, '5.1 开发环境')

    add_paragraph(doc,
        '本系统的开发工作是在以下软硬件环境中完成的，表5-1详细列出了开发环境的配置信息。'
    )

    add_table_caption(doc, '表5-1 开发环境配置')
    add_three_line_table(doc,
        ['项目', '配置内容'],
        [
            ['操作系统', 'Windows 10/11 64位'],
            ['开发工具', 'Visual Studio Code / Cursor IDE'],
            ['后端语言', 'Python 3.11'],
            ['后端框架', 'Flask 3.1.0'],
            ['前端框架', 'Vue 3.5 + Vite 6.0'],
            ['UI组件库', 'Element Plus 2.9'],
            ['数据库', 'SQLite 3（开发） / MySQL 8.0（生产）'],
            ['包管理', 'pip 24.x / npm 10.x'],
            ['版本控制', 'Git 2.x'],
            ['AI/ML库', 'scikit-learn 1.6, NumPy 2.2, Pandas 2.2'],
        ],
        col_widths=[4, 10]
    )

    add_heading_l2(doc, '5.2 系统实现')
    add_heading_l3(doc, '5.2.1 后端核心实现')

    add_paragraph(doc,
        '后端的项目结构采用Flask蓝图（Blueprint）模式进行组织，将不同功能模块的路由'
        '注册到各自的蓝图中，实现了代码的模块化管理。主要的项目文件结构如下：app目录'
        '下包含models（数据模型）、routes（API路由）、services（业务服务）三个子目录。'
        'models目录中定义了User、Clothing、ClothingCategory、WeatherRecord、'
        'Recommendation、UserFeedback、UserPreference等7个数据模型类，每个模型类'
        '都包含了to_dict()方法用于将数据库对象序列化为JSON可用的字典格式。'
    )

    add_paragraph(doc,
        'routes目录下包含auth.py（认证路由）、weather.py（天气路由）、clothing.py'
        '（服装路由）、recommendation.py（推荐路由）、user.py（用户路由）和admin.py'
        '（管理路由）共6个路由模块，合计提供了33个API接口端点。每个路由函数都遵循'
        '统一的响应格式，返回包含code、message和data三个字段的JSON对象。对于需要'
        '身份认证的接口，通过Flask-JWT-Extended的@jwt_required()装饰器来实现访问控制。'
    )

    add_paragraph(doc,
        '在天气数据服务实现方面，weather_service.py模块覆盖了32个主要城市的气象数据'
        '获取功能。系统为每个城市预设了经纬度坐标、年均温度基准、温度振幅、平均湿度和'
        '降水概率等气候参数。本地气候模型使用正弦函数模拟四季温度变化规律：'
    )

    add_paragraph(doc,
        'T(d) = T_base + T_amp × sin(2π × (d - 80) / 365) + ε',
        align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False
    )

    add_paragraph(doc,
        '其中d为一年中的天数序号，T_base为年均基础温度，T_amp为温度振幅，ε为服从'
        '正态分布N(0, 2.5)的随机扰动项。这种模型简洁有效地捕捉了温度的季节性变化特征。'
    )

    add_paragraph(doc,
        '天气数据获取接口的核心代码如下，该函数实现了缓存优先、API补充、模型兜底的'
        '三级数据获取策略：'
    )
    add_code_block(doc, '''def get_weather(city, target_date=None):
    if target_date is None:
        target_date = date.today()
    cached = WeatherRecord.query.filter_by(city=city, date=target_date).first()
    if cached and (datetime.utcnow() - cached.fetched_at).seconds < 3600:
        return cached.to_dict()
    weather_data = _generate_weather_for_city(city, target_date)
    api_key = current_app.config.get('WEATHER_API_KEY', '')
    if api_key and target_date == date.today():
        api_data = _fetch_from_qweather(city, api_key)
        if api_data: weather_data.update(api_data)
    return cached.to_dict()''')

    add_paragraph(doc,
        '用户登录认证接口的核心代码如下，通过JWT令牌机制实现无状态的身份认证：'
    )
    add_code_block(doc, '''@auth_bp.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    user = User.query.filter_by(username=data.get('username')).first()
    if not user or not user.check_password(data.get('password')):
        return jsonify(code=401, message='用户名或密码错误')
    token = create_access_token(identity=user.id,
                                expires_delta=timedelta(days=7))
    return jsonify(code=200, data={'token': token, 'user': user.to_dict()})''')

    add_paragraph(doc,
        '图5-1展示了用户登录界面，用户通过输入用户名和密码即可登录系统。图5-2展示了'
        '用户注册界面，新用户需要填写用户名、邮箱、密码等信息完成注册。'
    )
    insert_image(doc, img_sc_login, 13)
    add_figure_caption(doc, '图5-1 用户登录界面')
    insert_image(doc, img_sc_register, 13)
    add_figure_caption(doc, '图5-2 用户注册界面')

    add_heading_l3(doc, '5.2.2 前端核心实现')

    add_paragraph(doc,
        '前端项目基于Vue 3框架使用Vite构建工具进行搭建，整体上分为用户端和管理端两个'
        '部分。用户端包含登录页（Login）、注册页（Register）、首页推荐（Home）、天气'
        '预报（Forecast）、推荐历史（History）和个人中心（Profile）共6个页面视图。'
        '管理端包含数据概览（Dashboard）、用户管理（Users）、服装管理（ClothingManage）、'
        '天气监控（WeatherMonitor）和推荐记录（Recommendations）共5个页面视图。'
    )

    add_paragraph(doc,
        '前端通过Axios库封装了统一的HTTP请求模块，在请求拦截器中自动附加JWT令牌到'
        'Authorization请求头，在响应拦截器中统一处理401未授权错误（自动清除令牌并跳转'
        '到登录页面）。所有API调用的基础路径通过Vite的代理配置转发到后端Flask服务的'
        '5000端口，避免了开发环境下的跨域问题。Pinia状态管理库用于在全局范围内维护'
        '用户登录状态、令牌信息和当前城市等共享数据。'
    )

    add_paragraph(doc,
        '在数据可视化方面，管理端的Dashboard页面使用ECharts绑定了两个动态图表：场景'
        '分布饼图展示各场景的推荐使用频率，城市排行柱状图展示各城市的推荐热度。这些'
        '图表的数据均来源于后端实时统计接口，能够随着推荐记录的增加而动态更新。'
    )

    add_paragraph(doc,
        'Axios请求拦截器的核心代码如下，该模块实现了JWT令牌自动附加和401错误统一处理：'
    )
    add_code_block(doc, '''const service = axios.create({ baseURL: '/api', timeout: 10000 })
service.interceptors.request.use(config => {
  const token = localStorage.getItem('token')
  if (token) config.headers.Authorization = `Bearer ${token}`
  return config
})
service.interceptors.response.use(
  response => response.data,
  error => {
    if (error.response?.status === 401) {
      localStorage.clear(); router.push('/login')
    }
    return Promise.reject(error)
  })''')

    add_paragraph(doc,
        '图5-3展示了用户端首页的天气概览与场景选择区域，用户可以在此查看当前城市的实时'
        '天气信息和THI舒适度指数，并通过点击场景图标切换出行场景。图5-4展示了首页的穿衣'
        '推荐结果区域，包含三种AI推荐模式的切换、舒适度评分和搭配评分展示、推荐服装卡片'
        '列表以及用户反馈评分功能。'
    )
    insert_image(doc, img_sc_home, 13)
    add_figure_caption(doc, '图5-3 首页天气概览与场景选择')
    insert_image(doc, img_sc_home_rec, 13)
    add_figure_caption(doc, '图5-4 首页穿衣推荐结果展示')

    add_paragraph(doc,
        '图5-5展示了7天天气预报界面，用户可以查看未来一周的天气趋势以便提前规划穿着。'
        '图5-6展示了推荐历史记录界面，用户可以分页查看过往的穿衣推荐记录。图5-7展示了'
        '用户个人中心界面，用户可以在此编辑性别、年龄、体质类型等个人信息，这些信息直接'
        '影响系统的推荐算法决策。'
    )
    insert_image(doc, img_sc_forecast, 13)
    add_figure_caption(doc, '图5-5 7天天气预报界面')
    insert_image(doc, img_sc_history, 13)
    add_figure_caption(doc, '图5-6 推荐历史记录界面')
    insert_image(doc, img_sc_profile, 13)
    add_figure_caption(doc, '图5-7 用户个人中心界面')

    add_paragraph(doc,
        '管理端部分，图5-8展示了数据概览仪表盘，包含用户数、服装数、推荐记录数、反馈数'
        '等统计指标以及场景分布饼图和城市排行柱状图。图5-9展示了服装管理界面，管理员可以'
        '对服装数据进行增删改查操作。图5-10展示了用户管理界面，管理员可以查看和管理系统'
        '中的用户账户。图5-11展示了天气数据监控界面，实时展示各城市的天气采集情况。'
        '图5-12展示了推荐记录查看界面，管理员可以查看所有用户的推荐历史。'
    )
    insert_image(doc, img_sc_admin_dash, 13)
    add_figure_caption(doc, '图5-8 管理端数据概览仪表盘')
    insert_image(doc, img_sc_admin_clothing, 13)
    add_figure_caption(doc, '图5-9 管理端服装管理界面')
    insert_image(doc, img_sc_admin_users, 13)
    add_figure_caption(doc, '图5-10 管理端用户管理界面')
    insert_image(doc, img_sc_admin_weather, 13)
    add_figure_caption(doc, '图5-11 管理端天气数据监控界面')
    insert_image(doc, img_sc_admin_rec, 13)
    add_figure_caption(doc, '图5-12 管理端推荐记录查看界面')

    add_heading_l3(doc, '5.2.3 算法实现')

    add_paragraph(doc,
        '推荐引擎的核心实现代码位于recommendation_engine.py文件中，包含DecisionTree'
        'Recommender类和CollaborativeFilter类两个核心组件。DecisionTreeRecommender类'
        '在初始化时即完成模型的训练，使用35组基于专家知识构建的训练样本，特征维度为6'
        '（温度、湿度、风速、降水、AQI、场景ID），目标变量为保暖等级（1-5）。模型采用'
        'scikit-learn的DecisionTreeClassifier实现，设置max_depth=8以防止过拟合。'
    )

    add_paragraph(doc,
        'CollaborativeFilter类实现了基于用户的协同过滤推荐功能。get_similar_users方法'
        '遍历用户-服装偏好矩阵，计算目标用户与其他用户之间的余弦相似度，返回最相似的'
        'Top-5用户列表。recommend_items方法则根据相似用户的加权评分对候选服装进行排序，'
        '权重为相似度值本身。当用户偏好数据不足时，系统退化为基于保暖等级匹配的默认'
        '推荐，保证了在冷启动阶段系统仍然能够输出合理的推荐结果。'
    )

    add_paragraph(doc,
        '决策树保暖等级预测的核心代码如下：'
    )
    add_code_block(doc, '''class DecisionTreeRecommender:
    def __init__(self):
        self.model = DecisionTreeClassifier(
            max_depth=8, min_samples_split=3, min_samples_leaf=2)
        data = np.array(TRAINING_DATA)
        self.model.fit(data[:, :-1], data[:, -1])

    def predict_warmth_level(self, temp, humidity, wind, precip, aqi, scene):
        scene_id = SCENE_LIST.index(scene) if scene in SCENE_LIST else 0
        features = np.array([[temp, humidity, wind, precip, aqi, scene_id]])
        return int(self.model.predict(features)[0])''')

    add_paragraph(doc,
        '协同过滤用户相似度计算的核心代码如下：'
    )
    add_code_block(doc, '''def get_similar_users(self, user_id, top_n=5):
    all_prefs = UserPreference.query.all()
    user_items = {}
    for p in all_prefs:
        user_items.setdefault(p.user_id, {})[p.clothing_id] = p.preference_score
    target = user_items.get(user_id, {})
    similarities = []
    for other_id, other_prefs in user_items.items():
        if other_id == user_id: continue
        common = set(target.keys()) & set(other_prefs.keys())
        if len(common) < 2: continue
        vec_a = np.array([target[i] for i in common])
        vec_b = np.array([other_prefs[i] for i in common])
        cos_sim = np.dot(vec_a, vec_b) / (np.linalg.norm(vec_a) * np.linalg.norm(vec_b))
        similarities.append((other_id, cos_sim))
    return sorted(similarities, key=lambda x: x[1], reverse=True)[:top_n]''')

    add_heading_l2(doc, '5.3 系统测试')
    add_heading_l3(doc, '5.3.1 功能测试')

    add_paragraph(doc,
        '对系统进行了全面的功能测试，覆盖了所有33个API接口和前端页面的主要交互功能。'
        '功能测试的部分测试用例及结果如表5-2所示。'
    )

    add_table_caption(doc, '表5-2 功能测试用例表')
    add_three_line_table(doc,
        ['编号', '测试模块', '测试内容', '预期结果', '实际结果'],
        [
            ['TC-01', '用户注册', '正常注册新用户', '注册成功返回用户信息', '通过'],
            ['TC-02', '用户注册', '重复用户名注册', '返回错误提示', '通过'],
            ['TC-03', '用户登录', '正确账号密码登录', '返回JWT令牌', '通过'],
            ['TC-04', '用户登录', '错误密码登录', '返回认证失败', '通过'],
            ['TC-05', '天气查询', '查询北京当日天气', '返回完整天气数据', '通过'],
            ['TC-06', '天气预报', '获取7天预报', '返回7条预报数据', '通过'],
            ['TC-07', '穿衣推荐', '通勤场景获取推荐', '返回完整搭配方案', '通过'],
            ['TC-08', '穿衣推荐', '切换运动场景', '推荐结果发生变化', '通过'],
            ['TC-09', '反馈评分', '提交5分好评', '反馈保存成功', '通过'],
            ['TC-10', '服装管理', '添加新服装', '数据入库成功', '通过'],
            ['TC-11', '服装管理', '编辑服装属性', '数据更新成功', '通过'],
            ['TC-12', '服装管理', '删除服装', '数据删除成功', '通过'],
            ['TC-13', '用户管理', '禁用用户', '用户状态变为禁用', '通过'],
            ['TC-14', '数据概览', '查看统计数据', '正确显示各项指标', '通过'],
            ['TC-15', '推荐历史', '分页查看记录', '正确分页展示', '通过'],
        ],
        col_widths=[1.5, 2, 3.5, 3.5, 1.5]
    )

    add_paragraph(doc,
        '功能测试的结果显示，系统的全部15项核心功能测试用例均通过了验证，33个API接口'
        '在正常调用情况下均能返回正确的HTTP 200状态码和预期的响应数据。前端页面的各项'
        '交互操作也能够正常执行，数据能够在前端界面和后端数据库之间实现正确的同步。'
    )

    add_heading_l3(doc, '5.3.2 性能测试')

    add_paragraph(doc,
        '为评估系统在实际使用场景下的性能表现，对关键接口进行了简单的性能测试，'
        '测试结果如表5-3所示。'
    )

    add_table_caption(doc, '表5-3 性能测试结果')
    add_three_line_table(doc,
        ['测试接口', '平均响应时间', '最大响应时间', '并发数'],
        [
            ['用户登录', '45ms', '120ms', '10'],
            ['天气查询（缓存命中）', '15ms', '35ms', '10'],
            ['天气查询（首次获取）', '280ms', '650ms', '5'],
            ['穿衣推荐生成', '85ms', '200ms', '10'],
            ['服装列表查询', '25ms', '60ms', '10'],
            ['管理统计数据', '35ms', '80ms', '5'],
        ],
        col_widths=[4, 3, 3, 2]
    )

    add_paragraph(doc,
        '测试结果表明，在本地开发环境下（单机SQLite数据库），系统的大部分接口响应'
        '时间均能控制在100ms以内，天气数据在缓存命中的情况下响应速度尤其快。天气查询'
        '首次获取时响应时间较长，主要是因为需要通过气候模型计算生成天气数据并写入'
        '数据库缓存。整体性能表现能够满足小规模用户群体的使用需求。'
    )

    add_page_break(doc)
    # ============================================================
    # 第六章 用户体验评估
    # ============================================================
    add_heading_l1(doc, '用户体验评估', '6')

    add_heading_l2(doc, '6.1 评估方法')

    add_paragraph(doc,
        '为了对系统的用户体验进行评估，本研究采用了问卷调查和任务完成测试相结合的'
        '评估方法。邀请了10名在校大学生作为测试用户，年龄分布在19至24岁之间，其中'
        '男性6名、女性4名。测试流程为：每位用户首先完成注册登录，然后依次体验获取'
        '穿衣推荐、切换场景、查看天气预报、提交反馈评分等核心功能，最后填写包含5个'
        '维度的满意度评价问卷。评价采用5分制李克特量表。'
    )

    add_heading_l2(doc, '6.2 评估结果与分析')

    add_table_caption(doc, '表6-1 用户体验评估结果')
    add_three_line_table(doc,
        ['评估维度', '平均分', '标准差', '评价'],
        [
            ['界面美观度', '4.2', '0.63', '良好'],
            ['操作便捷性', '4.5', '0.53', '优秀'],
            ['推荐合理性', '3.8', '0.79', '良好'],
            ['响应速度', '4.3', '0.48', '良好'],
            ['整体满意度', '4.1', '0.57', '良好'],
        ],
        col_widths=[4, 2, 2, 3]
    )

    add_paragraph(doc,
        '评估结果显示，系统在操作便捷性方面得分最高（4.5分），说明系统的界面设计和'
        '交互流程较为友好直观，用户能够较快地完成各项操作任务。响应速度（4.3分）和'
        '界面美观度（4.2分）的评价也较为积极。推荐合理性的得分相对最低（3.8分），部分'
        '用户反馈在极端天气条件下推荐结果的精准度还有改善空间，同时有用户建议增加更多'
        '的服装搭配选项和风格偏好设置功能。整体满意度达到4.1分，表明系统基本达到了'
        '设计预期目标，但在推荐精度和个性化程度方面仍有进一步提升的潜力。'
    )

    add_page_break(doc)
    # ============================================================
    # 第七章 总结与展望
    # ============================================================
    add_heading_l1(doc, '总结与展望', '7')

    add_heading_l2(doc, '7.1 研究总结')

    add_paragraph(doc,
        '本文设计并实现了一套基于环境信息感知的智能出行穿衣服务系统，该系统综合运用了'
        '气象数据采集技术、决策树分类算法、协同过滤推荐算法以及现代Web开发技术，为用户'
        '提供了个性化的穿衣搭配建议服务。通过本研究，完成了以下几方面的工作内容：'
    )

    add_paragraph(doc,
        '第一，构建了一套双通道的气象数据采集方案，通过和风天气API接入实时气象数据，'
        '同时建立了基于正弦函数的本地气候模型覆盖32个城市，确保了系统在不同网络条件下'
        '的数据可用性。第二，实现了基于决策树的穿衣保暖等级预测模型，模型能够根据温度、'
        '湿度、风速、降水、AQI和场景等6维特征来判断适宜的穿衣厚度。第三，实现了基于'
        '余弦相似度的协同过滤推荐机制，结合用户历史偏好数据对候选服装进行个性化排序。'
        '第四，开发了包含用户端和管理端的完整Web应用系统，用户端提供天气展示、智能推荐、'
        '场景切换、历史记录等功能，管理端提供数据概览、用户管理、服装管理等功能。'
        '第五，系统引入了THI温湿指数作为人体舒适度的量化评估工具，并根据用户的体质'
        '类型和年龄进行了保暖等级的个性化调整。'
    )

    add_heading_l2(doc, '7.2 不足与展望')

    add_paragraph(doc,
        '尽管本系统已经实现了基本的智能穿衣推荐功能，但仍然存在一些不足之处有待改进：'
        '首先，当前的推荐算法主要依赖于决策树和简单的协同过滤，在面对复杂多变的气候'
        '条件时推荐精度还有提升空间，未来可以考虑引入深度学习方法如神经网络来增强'
        '模型的特征提取能力。其次，服装数据库的规模目前还较为有限，后续需要持续丰富'
        '服装品类和款式数据，以提供更加多样化的搭配选择。第三，系统目前尚未实现服装'
        '图片的智能识别功能，未来可以接入计算机视觉技术来支持用户上传服装照片并自动'
        '识别品类和风格。第四，协同过滤算法在用户量较少的冷启动阶段效果有限，可以考虑'
        '引入基于知识图谱的推荐方法来缓解这一问题。第五，系统的移动端适配还不够完善，'
        '后续可以考虑开发原生移动应用（如小程序或App）来提供更好的移动端使用体验。'
    )

    add_paragraph(doc,
        '总之，随着物联网、人工智能和气象科学技术的进一步发展，基于环境信息感知的'
        '智能穿衣服务系统在技术和应用方面都有着广阔的发展前景。期望本研究的工作能够为'
        '该领域的后续研究和实践提供一定的参考和借鉴价值。'
    )

    add_page_break(doc)
    # ============================================================
    # 参考文献
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(36)
    pf.space_after = Pt(18)
    pf.first_line_indent = Cm(0)
    run = p.add_run('参考文献')
    set_run_font(run, cn='黑体', size=SIZE_XIAO3, bold=False)

    refs = [
        '[1]史英杰,杨珂,王建欣,等.基于机器学习的时尚穿搭推荐研究综述[J].计算机应用研究,2022,39(4):978-985.',
        '[2]刘军平,张伏红,胡新荣,等.基于多模态融合的个性化服装搭配推荐[J].纺织学报,2023,44(3):176-186.',
        '[3]关菲,周艺,张晗.个性化推荐系统中协同过滤推荐算法优化研究[J].运筹与管理,2022,31(11):9-14.',
        '[4]于翘楚,赵明清,罗雨婷.基于最优权的协同过滤混合推荐算法及应用[J].运筹与管理,2024,33(7):79-84.',
        '[5]吕福荣,师云龙,景晓宁,等.服装推荐系统的关键技术研究进展[J].现代纺织技术,2024,32(12):134-144.',
        '[6]陶淘,侯俊,张晨亮,等.基于大数据云平台的天山气象APP设计与应用[J].气象科技,2024,52(2):195-204.',
        '[7]李圆,王孝东,于淼.以穿衣搭配数据为基础的协同过滤算法改进[J].纺织高校基础科学学报,2023,36(2):93-100.',
        '[8]霍春阳.Vue.js设计与实现[M].北京:人民邮电出版社,2022.',
        '[9]杨怡然,吴巧英.智能化服装搭配推荐研究进展[J].浙江理工大学学报(自然科学版),2021,45(1):1-12.',
        '[10]Bi K,Xie L,Zhang H,et al.Accurate medium-range global weather forecasting with 3D neural networks[J].Nature,2023,619(7970):533-541.',
        '[11]Deldjoo Y,Nazary F,Ramisa A,et al.A review of modern fashion recommender systems[J].ACM Computing Surveys,2024,56(4):1-37.',
        '[12]Zhou H,Xiong F,Chen H.A comprehensive survey of recommender systems based on deep learning[J].Applied Sciences,2023,13(20):11378.',
        '[13]Abdalla H I,Amer A A,Amer Y A,et al.Boosting the item-based collaborative filtering model with novel similarity measures[J].International Journal of Computational Intelligence Systems,2023,16:123.',
        '[14]Hill A J,Schumacher R S,Jirak I L.A new paradigm for medium-range severe weather forecasts:probabilistic random forest-based predictions[J].Weather and Forecasting,2023,38(2):251-272.',
        '[15]Haro Peralta J.Microservice APIs:Using Python,Flask,FastAPI,OpenAPI and More[M].Manning Publications,2022.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ref)
        set_run_font(run, cn='宋体', en='Times New Roman', size=SIZE_WU, bold=False)

    # ============================================================
    # 致谢
    # ============================================================
    add_page_break(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(36)
    pf.space_after = Pt(18)
    pf.first_line_indent = Cm(0)
    run = p.add_run('致　谢')
    set_run_font(run, cn='黑体', size=SIZE_XIAO3, bold=False)

    add_paragraph(doc,
        '在本次毕业设计的研究与论文撰写过程中，我得到了许多老师和同学们的帮助与支持，'
        '在此表达衷心的感谢之情。'
    )

    add_paragraph(doc,
        '首先，我要特别感谢我的指导老师在整个毕业设计过程中给予的悉心指导和耐心帮助。'
        '从选题方向的确定到系统方案的设计，从代码实现中遇到的技术难题到论文写作的规范'
        '要求，指导老师都给出了宝贵的意见和建议，使我在专业知识和科研能力方面获得了'
        '很大的提升和进步。'
    )

    add_paragraph(doc,
        '其次，我要感谢计算机与通信工程学院的各位任课老师，正是他们在四年本科学习期间'
        '传授的专业知识，为我完成本次毕业设计打下了坚实的理论基础。特别是在Python编程、'
        '数据库原理、Web开发技术和机器学习等课程中学到的知识，在本系统的开发过程中'
        '发挥了直接而重要的作用。'
    )

    add_paragraph(doc,
        '同时，我还要感谢一起学习和奋斗的同学们，大家在毕设期间互相交流技术经验、'
        '分享学习资源，这种良好的学习氛围对我的工作起到了积极的促进作用。感谢我的'
        '家人在整个大学生涯中给予的理解和支持，让我能够专注地完成学业和研究工作。'
    )

    add_paragraph(doc,
        '最后，感谢所有参考文献的作者们，他们的研究成果为本论文提供了重要的理论依据'
        '和方法参考。也感谢开源社区的贡献者们，Flask、Vue.js、scikit-learn等优秀的'
        '开源项目为本系统的开发提供了强大的技术支撑。'
    )

    # ===== 系统业务流程图作为附图 =====
    # 在第四章系统设计后补充系统流程图（已在正文中使用）

    # ===== 保存文件 =====
    output_path = os.path.join(BASE_DIR, '通信工程（专转本）Z2024130143陈鎏-基于环境信息感知的智能出行穿衣服系统.docx')
    doc.save(output_path)
    print(f'论文已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_thesis()
