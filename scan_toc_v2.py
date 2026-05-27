import docx
import re

doc = docx.Document(r'F:\26毕设2\基于环境信息感知的智能出行穿衣服务系统\通信工程（专转本）Z2024130143陈鎏-基于环境信.docx')

table_pattern = re.compile(r'^表\s*[\d\.]+\s+')
figure_pattern = re.compile(r'^图\s*[\d\.]+\s+')

print("=" * 80)
print("BODY TABLES (actual captions):")
print("=" * 80)
body_tables = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if not text:
        continue
    if style in ('table of figures',):
        continue
    if table_pattern.match(text):
        body_tables.append((i, text, style))
        print(f"  [{i:4d}] [{style:15s}] {text}")

print(f"\nTotal: {len(body_tables)} tables")

print("\n" + "=" * 80)
print("BODY FIGURES (actual captions):")
print("=" * 80)
body_figures = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if not text:
        continue
    if style in ('table of figures',):
        continue
    if figure_pattern.match(text):
        body_figures.append((i, text, style))
        print(f"  [{i:4d}] [{style:15s}] {text}")

print(f"\nTotal: {len(body_figures)} figures")

print("\n" + "=" * 80)
print("EXISTING TOC ENTRIES (style='table of figures'):")
print("=" * 80)
toc_entries = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    if style == 'table of figures' and text:
        toc_entries.append((i, text, style))
        print(f"  [{i:4d}] {text}")

print(f"\nTotal: {len(toc_entries)} TOC entries")

print("\n" + "=" * 80)
print("CONTEXT AROUND TOC (para 40-90):")
print("=" * 80)
for i in range(40, min(95, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ''
    if text:
        print(f"  [{i:4d}] [{style:20s}] {text[:100]}")
