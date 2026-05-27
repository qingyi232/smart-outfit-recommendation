# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

docx_path = r'F:\26毕设2\基于环境信息感知的智能出行穿衣服务系统\通信工程（专转本）Z2024130143陈鎏-毕业论文-修改版.docx'
output_path = r'F:\26毕设2\基于环境信息感知的智能出行穿衣服务系统\通信工程（专转本）Z2024130143陈鎏-毕业论文-修改版.docx'

doc = Document(docx_path)

formulas = {
    'THI = (1.8': ('2-1', 2),
    'sim(u, v) = cos': ('2-2', 2),
    'T(d) = T_base': ('5-1', 5),
}

chapter = 0
changes = 0

for para in doc.paragraphs:
    full_text = para.text.strip()
    
    ch_match = re.match(r'^(\d+)\s+\S', full_text)
    if ch_match and len(full_text) < 40:
        chapter = int(ch_match.group(1))
    
    for key, (num, ch) in formulas.items():
        if key in full_text and len(full_text) < 120:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 在段落末尾添加制表符和编号
            run = para.add_run('\t(' + num + ')')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置右对齐制表位
            pPr = para._element.get_or_add_pPr()
            tabs = pPr.makeelement(qn('w:tabs'), {})
            tab = tabs.makeelement(qn('w:tab'), {
                qn('w:val'): 'right',
                qn('w:pos'): '8306'  # 约14.7cm位置（A4页面右边距）
            })
            tabs.append(tab)
            pPr.append(tabs)
            
            changes += 1
            print(f'已为公式添加编号 ({num}): {full_text[:50]}...')

doc.save(output_path)
print(f'\n共处理 {changes} 个公式编号')
